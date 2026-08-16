# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document('تطوير نظام ادارة المستشفى - كامل.docx')
body = next(i for i,p in enumerate(doc.paragraphs) if p.text.strip()=='ملخص المشروع' and p.style.name=='main title')

for i in range(900, 1020):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if t:
        print(f'{i:4d}| {p.style.name:15s}| {t[:75]}')
