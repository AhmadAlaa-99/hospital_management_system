# -*- coding: utf-8 -*-
"""Pass 4: Fix chapter 5 numbering/styles and final TOC cleanup."""
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = 'تطوير نظام ادارة المستشفى - كامل.docx'

TOC_47 = [
    '4.7.6 الموقع الإلكتروني العام',
    '4.7.7 نظام الإشعارات والتواصل',
    '4.7.8 الوصفات الإلكترونية (e-Prescription)',
    '4.7.9 لوحة التقارير والإحصائيات',
    '4.7.10 جدولة أوقات الطبيب',
    '4.7.11 تكامل التأمين الكامل',
    '4.7.12 تذكير المواعيد',
    '4.7.13 تقييم الأطباء',
    '4.7.14 إدارة المدونة',
    '4.7.15 طلب الإسعاف',
    '4.7.16 تصدير السجل الطبي PDF',
    '4.7.17 إدارة الطابور (Queue Management)',
    '4.7.18 بيانات تجريبية (Database Seeders)',
]


def delete_paragraph(para):
    el = para._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def body_start(doc):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'ملخص المشروع' and p.style.name == 'main title':
            return i
    return 0


def insert_after(paragraph, text, style='arabic', bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def set_para(para, text, style='arabic', bold=False, center=False):
    para.style = style
    para.text = text
    if para.runs:
        para.runs[0].bold = bold
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def fix_toc_final(doc):
    toc_i = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == 'فهرس المحتويات')
    body_i = body_start(doc)

    # Remove junk / duplicate toc lines
    remove_texts = {
        'Summary and Conclusions', 'Future Prospects', 'Results and Conclusions',
        '5.2.5 اختبار الميزات المتقدمة',
    }
    to_remove = []
    for para in doc.paragraphs[toc_i:body_i]:
        t = para.text.strip()
        if t in remove_texts:
            to_remove.append(para)
        elif para.style.name.startswith('toc') and t.startswith('5.') and t not in {
            '5.1 أنواع الاختبارات', '5.2 سيناريوهات الاختبار',
            '5.3 نتائج الاختبار', '5.4 خاتمة الفصل',
        }:
            to_remove.append(para)
    for para in to_remove:
        delete_paragraph(para)

    to_remove = []
    for para in doc.paragraphs[toc_i:body_i]:
        t = para.text.strip()
        if para.style.name.startswith('toc') and re.match(r'^4\.7\.\d+', t):
            to_remove.append(para)
    for para in to_remove:
        delete_paragraph(para)

    parent = None
    for para in doc.paragraphs[toc_i:body_i]:
        t = para.text.strip()
        if t.startswith('4.7 ') or t == '4.7 تنفيذ الوحدات الرئيسية':
            parent = para
            break

    if parent:
        last = parent
        for title in TOC_47:
            last = insert_after(last, title, style='toc 3')

    # Ensure chapter 5-9 toc block
    anchor410 = None
    for para in doc.paragraphs[toc_i:body_i]:
        if para.text.strip().startswith('4.10'):
            anchor410 = para
    existing = {p.text.strip() for p in doc.paragraphs[toc_i:body_i]}
    if anchor410:
        block = [
            ('الفصل الخامس — الاختبار', 'toc 1'),
            ('5.1 أنواع الاختبارات', 'toc 3'),
            ('5.2 سيناريوهات الاختبار', 'toc 3'),
            ('5.3 نتائج الاختبار', 'toc 3'),
            ('5.4 خاتمة الفصل', 'toc 3'),
        ]
        last = anchor410
        for title, style in block:
            if title not in existing:
                last = insert_after(last, title, style=style)

    for para in doc.paragraphs[toc_i:body_i]:
        t = para.text.strip()
        if t == 'الفصل السادس — النتائج والاستنتاجات' and para.style.name == 'toc 3':
            set_para(para, t, style='toc 1')


def fix_chapter5_body(doc):
    body_i = body_start(doc)
    ch5 = None
    ch6 = None
    for i, para in enumerate(doc.paragraphs[body_i:], body_i):
        t = para.text.strip()
        if 'الفصل الخامس' in t and ch5 is None:
            ch5 = i
        if 'الفصل السادس' in t and para.style.name == 'main title':
            ch6 = i
            break
    if ch5 is None or ch6 is None:
        return

    # Chapter title
    p = doc.paragraphs[ch5]
    set_para(p, 'الفصل الخامس — الاختبار', style='main title', bold=True, center=True)

    # English subtitle
    if ch5 + 1 < len(doc.paragraphs):
        set_para(doc.paragraphs[ch5 + 1], 'Testing', style='English', center=True)

    mapping = [
        (re.compile(r'^4\.1\b'), '5.1'),
        (re.compile(r'^4\.1\.1\b'), '5.1.1'),
        (re.compile(r'^4\.1\.2\b'), '5.1.2'),
        (re.compile(r'^4\.1\.3\b'), '5.1.3'),
        (re.compile(r'^4\.1\.4\b'), '5.1.4'),
        (re.compile(r'^4\.2\b'), '5.2'),
        (re.compile(r'^4\.2\.1\b'), '5.2.1'),
        (re.compile(r'^4\.2\.2\b'), '5.2.2'),
        (re.compile(r'^4\.2\.3\b'), '5.2.3'),
        (re.compile(r'^4\.2\.4\b'), '5.2.4'),
        (re.compile(r'^4\.3\b'), '5.3'),
        (re.compile(r'^4\.3\.1\b'), '5.3.1'),
        (re.compile(r'^4\.3\.2\b'), '5.3.2'),
        (re.compile(r'^4\.3\.3\b'), '5.3.3'),
        (re.compile(r'^4\.4\b'), '5.4'),
    ]

    for para in doc.paragraphs[ch5:ch6]:
        t = para.text.strip()
        if not t:
            continue
        new_t = t
        for rx, repl in mapping:
            new_t = rx.sub(repl, new_t, count=1)
        if new_t != t:
            # style by level
            if re.match(r'^5\.\d+\.\d+', new_t):
                st = 'Sub-Side title'
            elif re.match(r'^5\.\d+\s', new_t) or re.match(r'^5\.\d+$', new_t):
                st = 'Side title'
            else:
                st = para.style.name
            set_para(para, new_t, style=st, bold=True)

        # body lines wrongly styled
        if para.style.name == 'Sub-Side title' and (t.startswith('تم') or t.startswith('-') or t.startswith('في')):
            set_para(para, t, style='arabic')

    # Move 5.2.5 block after 5.2.4 content
    sec25 = None
    sec24_end = None
    for para in doc.paragraphs[ch5:ch6]:
        t = para.text.strip()
        if t.startswith('5.2.4'):
            sec24_end = para
        if t.startswith('5.2.5'):
            sec25 = para
    if sec25 and sec24_end:
        block = [sec25]
        nxt = sec25._p.getnext()
        while nxt is not None:
            p = Paragraph(nxt, sec25._parent)
            if p.text.strip().startswith('5.3') or p.text.strip().startswith('-  طلب فحص'):
                break
            if p.text.strip().startswith('-'):
                block.append(p)
                nxt = nxt.getnext()
            else:
                break
        # if 5.2.5 is before 4.2.4 tail, leave for manual - skip complex move


def main():
    doc = Document(SOURCE)
    fix_toc_final(doc)
    fix_chapter5_body(doc)
    doc.save(SOURCE)
    print('Pass 4 done.')


if __name__ == '__main__':
    main()
