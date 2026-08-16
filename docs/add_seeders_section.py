# -*- coding: utf-8 -*-
"""Add Database Seeders section to graduation Word document."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_before(paragraph, text, bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def insert_after(paragraph, text, bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def find_paragraph(doc, startswith=None, contains=None):
    for para in doc.paragraphs:
        t = para.text.strip()
        if startswith and t.startswith(startswith):
            return para
        if contains and contains in t:
            return para
    return None


def main():
    path = 'docs/تطوير نظام ادارة المستشفى - كامل.docx'
    doc = Document(path)

    if find_paragraph(doc, startswith='12.7.18'):
        print('Section 12.7.18 already exists')
        doc.save(path)
        doc.save('تطوير نظام ادارة المستشفى.docx')
        return

    anchor = find_paragraph(doc, startswith='12.8')
    if not anchor:
        anchor = find_paragraph(doc, startswith='12.7.17')
    if not anchor:
        print('Anchor not found')
        return

    blocks = [
        ('12.7.18 بيانات تجريبية (Database Seeders)', True),
        ('Seeder رئيسي: HmsFullDemoSeeder — يُحمّل كل بيانات العرض التجريبية.', False),
        ('الأمر: php artisan db:seed  أو  php artisan db:seed --class=HmsFullDemoSeeder', False),
        ('أو إعادة بناء كاملة: php artisan migrate:fresh --seed', False),
        ('', False),
        ('جدول Seeders:', True),
        ('| Seeder | المحتوى |', False),
        ('| UserTableSeeder | مستخدمون أساسيون |', False),
        ('| AdminTableSeeder | حساب المدير (admin@gmail.com) |', False),
        ('| SectionTableSeeder | 6 أقسام طبية |', False),
        ('| DoctorTableSeeder | أطباء + صور |', False),
        ('| PatientTableSeeder | 8 مرضى (patient@yahoo.com) |', False),
        ('| RayEmployeeTableSeeder / LaboratorieEmployeeTableSeeder | موظفو الأشعة والمختبر |', False),
        ('| ServiceTableSeeder / GroupTableSeeder | خدمات فردية ومجمّعة |', False),
        ('| AmbulanceInsuranceSeeder | 3 سيارات إسعاف + 3 شركات تأمين |', False),
        ('| AppointmentBookingSeeder | 12 موعد (معلق / مؤكد / منتهي) |', False),
        ('| InvoiceDemoSeeder | فواتير + تشخيصات + أشعة + مختبر + سندات |', False),
        ('| ExtendedFeaturesSeeder | جداول أطباء، وصفات، مطالبات تأمين، تقييمات، طلبات إسعاف |', False),
        ('| BlogSiteSettingSeeder | 4 مقالات + إعدادات الموقع |', False),
        ('| QueueTicketSeeder | أرقام طابور لكل قسم (انتظار / نداء / عند الطبيب / مكتمل) |', False),
        ('', False),
        ('حسابات تجريبية بعد التشغيل:', True),
        ('• المدير: admin@gmail.com / 123456789', False),
        ('• المريض: patient@yahoo.com / 12345678', False),
        ('• شاشة الطابور: /queue/display/section/1', False),
        ('• تتبع الانتظار: /queue/track', False),
    ]

    for text, bold in reversed(blocks):
        if text:
            insert_before(anchor, text, bold=bold)
        else:
            insert_before(anchor, ' ', bold=False)

    p131 = find_paragraph(doc, contains='15. إدارة الطابور')
    if p131:
        insert_after(p131, '16. بيانات Seeders تجريبية (HmsFullDemoSeeder)', False)

    doc.save(path)
    doc.save('تطوير نظام ادارة المستشفى.docx')
    print('Updated seeders section in Word doc')


if __name__ == '__main__':
    main()
