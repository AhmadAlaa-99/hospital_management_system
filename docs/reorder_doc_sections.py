# -*- coding: utf-8 -*-
"""Reorder 12.7.6-12.7.16 in ascending order before 12.8."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_before(paragraph, text, bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


BLOCKS = [
    ('12.7.6 الموقع الإلكتروني العام', True),
    ('صفحة رئيسية، حجز مواعيد، مدونة، طلب إسعاف، إعدادات الموقع.', False),
    ('12.7.7 نظام الإشعارات والتواصل', True),
    ('NotificationService، بريد تأكيد، SMS Twilio، إشعارات فورية.', False),
    ('12.7.8 الوصفات الإلكترونية (e-Prescription)', True),
    ('جدول prescriptions: دواء، جرعة، تكرار، مدة، تعليمات — مرتبط بالتشخيص.', False),
    ('12.7.9 لوحة التقارير والإحصائيات', True),
    ('/reports — Charts: مرضى/شهر، إيرادات، أداء أقسام، متوسط تقييم الأطباء.', False),
    ('12.7.10 جدولة أوقات الطبيب', True),
    ('doctor_schedules + API slots + منع التعارض.', False),
    ('12.7.11 تكامل التأمين الكامل', True),
    ('insurance_id، خصم تلقائي، مطالبات، تقارير الشركات.', False),
    ('12.7.12 تذكير المواعيد', True),
    ('appointments:send-reminders — Email + SMS قبل 24 ساعة.', False),
    ('12.7.13 تقييم الأطباء', True),
    ('doctor_ratings: نجوم 1–5 وتعليق بعد موعد منتهي.', False),
    ('12.7.14 إدارة المدونة', True),
    ('CRUD من Dashboard: /admin/blogs.', False),
    ('12.7.15 طلب الإسعاف', True),
    ('نموذج الموقع → ambulance_requests → إرسال سيارة.', False),
    ('12.7.16 تصدير السجل الطبي PDF', True),
    ('DomPDF — تشخيصات + وصفات + أشعة + مختبر — للطبيب والمريض والمدير.', False),
]


def main():
    doc = Document('تطوير نظام ادارة المستشفى.docx')
    paras = list(doc.paragraphs)

    anchor_128 = None
    after_275 = False
    to_remove = []
    for para in paras:
        t = para.text.strip()
        if t.startswith('12.7.5'):
            after_275 = True
            continue
        if after_275:
            if t.startswith('12.7.6') or t.startswith('12.7.7') or t.startswith('12.7.8') or \
               t.startswith('12.7.9') or t.startswith('12.7.10') or t.startswith('12.7.11') or \
               t.startswith('12.7.12') or t.startswith('12.7.13') or t.startswith('12.7.14') or \
               t.startswith('12.7.15') or t.startswith('12.7.16'):
                to_remove.append(para)
            elif any(k in t for k in ['NotificationService', 'DomPDF', 'doctor_schedules',
                                       'insurance_claims', 'appointments:send-reminders',
                                       'doctor_ratings', '/admin/blogs', 'ambulance_requests',
                                       'صفحة رئيسية', 'جدول prescriptions', '/reports']):
                to_remove.append(para)
            elif t.startswith('12.8'):
                anchor_128 = para
                break

    for para in to_remove:
        delete_paragraph(para)

    if anchor_128:
        for text, bold in reversed(BLOCKS):
            insert_before(anchor_128, text, bold=bold)

    doc.save('تطوير نظام ادارة المستشفى.docx')
    doc.save('docs/تطوير نظام ادارة المستشفى - كامل.docx')
    print('Reordered OK')


if __name__ == '__main__':
    main()
