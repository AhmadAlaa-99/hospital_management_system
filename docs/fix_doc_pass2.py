# -*- coding: utf-8 -*-
"""Second pass: fix misplaced TOC/body content after format_university_doc.py"""
from __future__ import annotations

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = 'تطوير نظام ادارة المستشفى - كامل.docx'

SECTIONS_476_418 = [
    ('4.7.6 الموقع الإلكتروني العام', 'Side title', True),
    ('صفحة رئيسية (Hero، إحصائيات، أقسام، أطباء، مدونة)، حجز مواعيد AJAX، طلب إسعاف، إعدادات الموقع.', 'arabic', False),
    ('4.7.7 نظام الإشعارات والتواصل', 'Side title', True),
    ('NotificationService للإشعارات الداخلية، بريد تأكيد الموعد (AppointmentConfirmation)، SMS عبر Twilio (اختياري).', 'arabic', False),
    ('4.7.8 الوصفات الإلكترونية (e-Prescription)', 'Side title', True),
    ('جدول prescriptions مرتبط بالتشخيص: اسم الدواء، الجرعة، التكرار، المدة (أيام)، التعليمات.', 'arabic', False),
    ('4.7.9 لوحة التقارير والإحصائيات', 'Side title', True),
    ('مسار /reports للمدير — رسوم Chart.js: مرضى جدد شهرياً، إيرادات، أداء الأقسام، متوسط تقييم الأطباء.', 'arabic', False),
    ('4.7.10 جدولة أوقات الطبيب', 'Side title', True),
    ('جدول doctor_schedules (يوم، من، إلى، مدة الموعد). AppointmentScheduleService يمنع تعارض المواعيد.', 'arabic', False),
    ('4.7.11 تكامل التأمين الكامل', 'Side title', True),
    ('ربط المريض بشركة تأمين (insurance_id)، خصم تلقائي، مطالبات insurance_claims، تقرير حسب شركة التأمين.', 'arabic', False),
    ('4.7.12 تذكير المواعيد', 'Side title', True),
    ('أمر appointments:send-reminders مجدول كل ساعة — Email + SMS قبل 24 ساعة.', 'arabic', False),
    ('4.7.13 تقييم الأطباء', 'Side title', True),
    ('جدول doctor_ratings: تقييم 1–5 وتعليق بعد موعد منتهي (مرة واحدة).', 'arabic', False),
    ('4.7.14 إدارة المدونة', 'Side title', True),
    ('CRUD من Dashboard (/admin/blogs). إعجاب وتعليق للمريض المسجّل؛ مودال دخول/تسجيل للزائر.', 'arabic', False),
    ('4.7.15 طلب الإسعاف', 'Side title', True),
    ('نموذج في الموقع → ambulance_requests → إرسال سيارة / إكمال / إلغاء.', 'arabic', False),
    ('4.7.16 تصدير السجل الطبي PDF', 'Side title', True),
    ('مكتبة DomPDF — تشخيصات + وصفات + أشعة + مختبر — للطبيب والمريض والمدير.', 'arabic', False),
    ('4.7.17 إدارة الطابور (Queue Management)', 'Side title', True),
    ('جدول queue_tickets — استقبال (/queue)، طبيب (/doctor/queue)، شاشة TV، تتبع (/queue/track)، Pusher.', 'arabic', False),
    ('4.7.18 بيانات تجريبية (Database Seeders)', 'Side title', True),
    ('Seeder رئيسي HmsFullDemoSeeder — php artisan db:seed | migrate:fresh --seed', 'arabic', False),
    ('حسابات تجريبية: admin@gmail.com / 123456789 — patient@yahoo.com / 12345678', 'arabic', False),
]

TOC_47 = [
    '4.7.6 الموقع الإلكتروني العام',
    '4.7.7 نظام الإشعارات والتواصل',
    '4.7.8 الوصفات الإلكترونية (e-Prescription)',
    '4.7.9 لوحة التقارير والإحصائيات',
    '4.7.10 جدولة أوقات الطبيب',
    '4.7.11 تكامل التأمين الكامل',
    '4.7.12 تذكير المواعيد',
    '4.7.13 تقييم الأطباء',
    '4.7.14 إدارة المدونة',
    '4.7.15 طلب الإسعاف',
    '4.7.16 تصدير السجل الطبي PDF',
    '4.7.17 إدارة الطابور (Queue Management)',
    '4.7.18 بيانات تجريبية (Database Seeders)',
]


def delete_paragraph(para):
    para._element.getparent().remove(para._element)


def body_index(doc):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'ملخص المشروع' and p.style.name == 'main title':
            return i
    # fallback: last occurrence
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'ملخص المشروع':
            idx = i
    return idx or 0


def find_body_paragraph(doc, **kw):
    start = body_index(doc)
    for para in doc.paragraphs[start:]:
        if para.style.name.startswith('toc'):
            continue
        t = para.text.strip()
        if kw.get('equals') and t == kw['equals']:
            return para
        if kw.get('startswith') and t.startswith(kw['startswith']):
            return para
        if kw.get('contains') and kw['contains'] in t:
            return para
    return None


def insert_before(paragraph, text, style='arabic', bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def set_para(para, text, style='arabic', bold=False):
    para.style = style
    para.text = text
    if bold and para.runs:
        para.runs[0].bold = True


def clean_toc_area(doc):
    toc_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == 'فهرس المحتويات')
    body_start = body_index(doc)

    # Remove non-toc content between TOC and body (including misplaced section bodies)
    for para in list(doc.paragraphs)[toc_start + 1:body_start]:
        if not para.style.name.startswith('toc') and para.text.strip():
            delete_paragraph(para)
        elif para.style.name == 'main title' and 'الفصل الخامس' in para.text:
            delete_paragraph(para)
        elif para.style.name in ('Side title', 'Sub-Side title', 'English') and para.text.strip():
            delete_paragraph(para)

    # Remove duplicate 4.1
    seen = set()
    for para in list(doc.paragraphs)[toc_start:body_start]:
        t = para.text.strip()
        if t in seen and para.style.name.startswith('toc'):
            delete_paragraph(para)
        else:
            seen.add(t)

    # Rebuild 4.7.x toc entries in order
    anchor = None
    for para in doc.paragraphs[toc_start:body_start]:
        if para.text.strip().startswith('4.8'):
            anchor = para
            break
    if anchor:
        # remove existing 4.7.x toc lines
        for para in list(doc.paragraphs)[toc_start:body_start]:
            t = para.text.strip()
            if para.style.name.startswith('toc') and t.startswith('4.7.'):
                delete_paragraph(para)
        for title in reversed(TOC_47):
            insert_before(anchor, title, style='toc 3')


def para_index(doc, target):
    for i, p in enumerate(doc.paragraphs):
        if p._p is target._p:
            return i
    return None


def remove_body_duplicates(doc):
    start = body_index(doc)
    anchor = find_body_paragraph(doc, startswith='4.8')
    if not anchor:
        return

    ai = para_index(doc, anchor)
    if ai is None:
        return

    to_del = []
    for para in doc.paragraphs[start:ai]:
        t = para.text.strip()
        if re.match(r'^4\.7\.([6-9]|1[0-8])', t):
            to_del.append(para)
        elif any(k in t for k in [
            'NotificationService', 'DomPDF', 'doctor_schedules', 'insurance_claims',
            'appointments:send-reminders', 'doctor_ratings', '/admin/blogs',
            'ambulance_requests', 'queue_tickets', 'HmsFullDemoSeeder', 'admin@gmail.com',
            'blog_likes', 'مودال', 'Hero،',
        ]):
            to_del.append(para)
    for para in to_del:
        delete_paragraph(para)

    anchor = find_body_paragraph(doc, startswith='4.8')
    if anchor:
        for text, style, bold in reversed(SECTIONS_476_418):
            insert_before(anchor, text, style=style, bold=bold)


def fix_chapter5_in_toc(doc):
    """Move chapter 5 content out of TOC if misplaced."""
    toc_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == 'فهرس المحتويات')
    body_start = body_index(doc)

    toc_ch5_titles = {
        'الفصل الخامس — الاختبار': 'toc 1',
        '5.1 أنواع الاختبارات': 'toc 3',
        '5.1.1 اختبار الوحدة (Unit Testing)': 'toc 3',
        '5.1.2 اختبار التكامل (Integration Testing)': 'toc 3',
        '5.1.3 اختبار النظام (System Testing)': 'toc 3',
        '5.1.4 اختبار القبول (Acceptance Testing)': 'toc 3',
        '5.2 سيناريوهات الاختبار': 'toc 3',
        '5.3 نتائج الاختبار': 'toc 3',
        '5.4 خاتمة الفصل': 'toc 3',
        'الفصل السادس — النتائج والاستنتاجات': 'toc 1',
        '6.1 النتائج الرئيسية': 'toc 3',
        '6.1.1 النتائج الوظيفية': 'toc 3',
        '6.1.2 النتائج التقنية': 'toc 3',
        '6.2 الإنجازات': 'toc 3',
        '6.3 الاستنتاجات': 'toc 3',
        '6.4 التوصيات': 'toc 3',
        '6.4.1 توصيات للتطوير المستقبلي': 'toc 3',
        '6.4.2 توصيات للاستخدام': 'toc 3',
        '6.5 خاتمة الفصل': 'toc 3',
        'الفصل السابع — الخلاصة والاستنتاجات': 'toc 1',
        'الفصل الثامن — الأفاق المستقبلية': 'toc 1',
        'الفصل التاسع — المراجع': 'toc 1',
    }

    for para in list(doc.paragraphs)[toc_start + 1:body_start]:
        t = para.text.strip()
        if t == 'Testing':
            delete_paragraph(para)
        elif t.startswith('تم اختبار') or t.startswith('- Models'):
            delete_paragraph(para)
        elif para.style.name == 'main title' and 'الفصل الخامس' in t:
            delete_paragraph(para)
        elif para.style.name in ('Side title', 'Sub-Side title', 'English') and t in toc_ch5_titles:
            delete_paragraph(para)

    # Ensure TOC has chapter 5-9 entries
    anchor = find_body_paragraph(doc, startswith='4.10') or find_body_paragraph(doc, startswith='4.8')
    # find last toc entry before body
    last_toc = None
    for para in doc.paragraphs[toc_start:body_start]:
        if para.style.name.startswith('toc'):
            last_toc = para
    if last_toc is None:
        return

    existing = {p.text.strip() for p in doc.paragraphs[toc_start:body_start]}
    insert_point = last_toc
    for title, style in [
        ('الفصل الخامس — الاختبار', 'toc 1'),
        ('5.1 أنواع الاختبارات', 'toc 3'),
        ('5.2 سيناريوهات الاختبار', 'toc 3'),
        ('5.3 نتائج الاختبار', 'toc 3'),
        ('5.4 خاتمة الفصل', 'toc 3'),
        ('الفصل السادس — النتائج والاستنتاجات', 'toc 1'),
        ('6.1 النتائج الرئيسية', 'toc 3'),
        ('6.2 الإنجازات', 'toc 3'),
        ('6.3 الاستنتاجات', 'toc 3'),
        ('6.4 التوصيات', 'toc 3'),
        ('6.5 خاتمة الفصل', 'toc 3'),
        ('الفصل السابع — الخلاصة والاستنتاجات', 'toc 1'),
        ('الفصل الثامن — الأفاق المستقبلية', 'toc 1'),
        ('الفصل التاسع — المراجع', 'toc 1'),
    ]:
        if title not in existing:
            insert_point = insert_before(insert_point, title, style=style)


def fix_misc(doc):
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith('12.4'):
            set_para(para, t.replace('12.4', '4.4.1', 1), style=para.style.name, bold=True)
        if para.style.name == 'Normal' and re.match(r'^\d+\.\s', t) and '6.1' not in t:
            # numbered results list items
            if any(k in t for k in ['الموقع', 'الوصفات', 'التقارير', 'جدولة', 'التأمين', 'تذكير', 'تقييم', 'المدونة', 'PDF', 'الطابور', 'Seeders']):
                set_para(para, t, style='arabic')

    # Ensure chapter 6 header exists
    ch6 = find_body_paragraph(doc, contains='الفصل السادس')
    if ch6 and ch6.style.name != 'main title':
        set_para(ch6, 'الفصل السادس — النتائج والاستنتاجات', style='main title', bold=True)
        ch6.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document(SOURCE)
    clean_toc_area(doc)
    remove_body_duplicates(doc)
    fix_chapter5_in_toc(doc)
    fix_misc(doc)
    doc.save(SOURCE)
    print('Second pass complete.')


if __name__ == '__main__':
    main()
