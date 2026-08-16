# -*- coding: utf-8 -*-
"""Capture HMS screenshots and insert them into the graduation Word document."""
from __future__ import annotations

import shutil
import time
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from playwright.sync_api import TimeoutError as PlaywrightTimeout
from playwright.sync_api import sync_playwright

BASE = "http://127.0.0.1:8000"
DOCS = Path(__file__).resolve().parent
SHOTS = DOCS / "screenshots" / "doc"
DOCX = DOCS / "تطوير نظام ادارة المستشfى - كامل.docx"
BACKUP = DOCS / f"backup_{datetime.now():%Y%m%d_%H%M%S}_لقطات.docx"

PASS = "12345678"
ACCOUNTS = {
    "admin": ("admin", "admin@gmail.com", PASS),
    "doctor": ("doctor", "doctor@gmail.com", PASS),
    "patient": ("user", "patient@yahoo.com", PASS),
    "ray_employee": ("ray_employee", "ray@hospital.com", PASS),
    "laboratorie_employee": ("laboratorie_employee", "lab@hospital.com", PASS),
}

HEADING_IMAGES = {
    "4.11.1.1 الصفحة الرئيسية": "4-11-1-1-homepage.png",
    "4.11.1.2 صفحة حجز موعد": "4-11-1-2-appointment.png",
    "4.11.1.3 صفحة المدونة (قائمة المقالات)": "4-11-1-3-blogs.png",
    "4.11.1.4 صفحة تفاصيل مقال (إعجاب/تعليق)": "4-11-1-4-blog-details.png",
    "4.11.1.5 صفحة طلب الإسعاف": "4-11-1-5-ambulance.png",
    "4.11.1.6 صفحة تتبع الانتظار (/queue/track)": "4-11-1-6-queue-track.png",
    "4.11.1.7 تسجيل/دخول المريض من الموقع": "4-11-1-7-patient-auth.png",
    "4.11.2.1 لوحة المعلومات الرئيسية (Dashboard)": "4-11-2-1-admin-dashboard.png",
    "4.11.2.2 إدارة الأقسام": "4-11-2-2-sections.png",
    "4.11.2.3 إدارة الأطباء": "4-11-2-3-doctors.png",
    "4.11.2.4 إدارة المرضى": "4-11-2-4-patients.png",
    "4.11.2.5 إدارة المواعيد": "4-11-2-5-appointments.png",
    "4.11.2.6 لوحة التقارير والإحصائيات (/reports)": "4-11-2-6-reports.png",
    "4.11.2.7 إدارة مطالبات التأمين": "4-11-2-7-claims.png",
    "4.11.2.8 إدارة طلبات الإسعاف": "4-11-2-8-ambulance-req.png",
    "4.11.2.9 إدارة المدونة (CRUD)": "4-11-2-9-blogs-admin.png",
    "4.11.2.10 إعدادات الموقع": "4-11-2-10-settings.png",
    "4.11.2.11 إدارة الانتظار — لوحة الاستقبال (/queue)": "4-11-2-11-queue.png",
    "4.11.2.12 جدولة أوقات الأطباء": "4-11-2-12-schedules.png",
    "4.11.3.1 مواعيد الطبيب (قائمة/منتهية)": "4-11-3-1-doctor-appointments.png",
    "4.11.3.2 قائمة انتظار العيادة (/doctor/queue)": "4-11-3-2-doctor-queue.png",
    "4.11.3.3 إضافة تشخيص ووصفة إلكترونية": "4-11-3-3-diagnosis.png",
    "4.11.3.4 طلب فحص أشعة/مختبر": "4-11-3-4-labs-rays.png",
    "4.11.3.5 تفاصيل المريض والسجل الطبي": "4-11-3-5-patient-details.png",
    "4.11.3.6 المحادثة مع المريض (Livewire Chat)": "4-11-3-6-chat.png",
    "4.11.4.1 مواعيدي": "4-11-4-1-my-appointments.png",
    "4.11.4.2 تقييم الطبيب بعد الموعد": "4-11-4-2-rating.png",
    "4.11.4.3 عرض نتائج الأشعة والمختبر": "4-11-4-3-results.png",
    "4.11.4.4 تصدير السجل الطبي PDF": "4-11-4-4-pdf.png",
    "4.11.4.5 الفواتير والمدفوعات": "4-11-4-5-invoices.png",
    "4.11.5.1 قائمة طلبات الأشعة": "4-11-5-1-ray-list.png",
    "4.11.5.2 إدخال نتيجة الأشعة": "4-11-5-2-ray-edit.png",
    "4.11.6.1 قائمة طلبات المختبر": "4-11-6-1-lab-list.png",
    "4.11.6.2 إدخال نتيجة المختبر": "4-11-6-2-lab-edit.png",
    "4.11.7.1 شاشة عرض الانتظار لقسم (/queue/display/section/{id})": "4-11-7-1-queue-tv.png",
    "4.11.7.2 تبديل الأقسام على شاشة العرض": "4-11-7-2-queue-tv-switch.png",
    "4.11.8.1 صفحة تسجيل الدخول (لوحات التحكم)": "4-11-8-1-login.png",
    "4.11.8.2 الإشعارات الداخلية": "4-11-8-2-notifications.png",
}


def clean_ui(page) -> None:
    page.evaluate(
        """() => {
            document.querySelectorAll('.preloader, .hms-global-loader').forEach(el => el.remove());
            const loaders = document.querySelectorAll('[class*="loader"]');
            loaders.forEach(el => { if (el.classList.contains('preloader')) el.remove(); });
        }"""
    )


def shot(page, name: str, full: bool = False) -> None:
    SHOTS.mkdir(parents=True, exist_ok=True)
    clean_ui(page)
    page.wait_for_timeout(400)
    page.screenshot(path=str(SHOTS / name), full_page=full)
    print("saved", name)


def login(page, role: str) -> None:
    panel, email, password = ACCOUNTS[role]
    page.context.clear_cookies()
    page.goto(f"{BASE}/login", wait_until="domcontentloaded", timeout=30000)
    page.wait_for_selector("#sectionChooser", timeout=20000)
    clean_ui(page)
    page.select_option("#sectionChooser", panel)
    page.wait_for_timeout(400)
    page.locator(f"#{panel} input[name='email']").fill(email)
    page.locator(f"#{panel} input[name='password']").fill(password)
    page.locator(f"#{panel} button[type='submit']").click()
    page.wait_for_load_state("domcontentloaded")
    page.wait_for_timeout(800)
    clean_ui(page)


def open_public(page, path: str) -> None:
    page.goto(f"{BASE}{path}", wait_until="domcontentloaded", timeout=30000)
    clean_ui(page)
    page.wait_for_timeout(600)


def capture(page) -> None:
    # Public website
    open_public(page, "/")
    shot(page, "4-11-1-1-homepage.png")

    try:
        page.locator("#appointment").scroll_into_view_if_needed()
        page.wait_for_timeout(500)
        shot(page, "4-11-1-2-appointment.png")
    except Exception as exc:
        print("appointment scroll failed", exc)

    try:
        page.locator(".hms-ambulance-section, #contact").first.scroll_into_view_if_needed()
        page.wait_for_timeout(500)
        shot(page, "4-11-1-5-ambulance.png")
    except Exception as exc:
        print("ambulance scroll failed", exc)

    open_public(page, "/blogs")
    shot(page, "4-11-1-3-blogs.png")

    open_public(page, "/blogs/children-winter-care-tips")
    shot(page, "4-11-1-4-blog-details.png")

    open_public(page, "/queue/track")
    shot(page, "4-11-1-6-queue-track.png")

    open_public(page, "/")
    try:
        page.evaluate("document.body.classList.add('hms-auth-open'); document.querySelector('.hms-auth-modal')?.classList.add('is-open');")
        page.wait_for_timeout(400)
        shot(page, "4-11-1-7-patient-auth.png")
    except Exception as exc:
        print("auth modal failed", exc)

    open_public(page, "/queue/display/section/7")
    shot(page, "4-11-7-1-queue-tv.png")
    try:
        page.locator("text=NEUR").first.click(timeout=3000)
        page.wait_for_timeout(700)
    except PlaywrightTimeout:
        try:
            page.locator("a,button").filter(has_text="المخ").first.click(timeout=3000)
            page.wait_for_timeout(700)
        except Exception:
            pass
    shot(page, "4-11-7-2-queue-tv-switch.png")

    open_public(page, "/login")
    shot(page, "4-11-8-1-login.png")

    # Admin
    login(page, "admin")
    page.goto(f"{BASE}/dashboard/admin", wait_until="domcontentloaded")
    clean_ui(page)
    shot(page, "4-11-2-1-admin-dashboard.png")

    for path, name in [
        ("/Sections", "4-11-2-2-sections.png"),
        ("/Doctors", "4-11-2-3-doctors.png"),
        ("/Patients", "4-11-2-4-patients.png"),
        ("/appointments", "4-11-2-5-appointments.png"),
        ("/reports", "4-11-2-6-reports.png"),
        ("/insurance-claims", "4-11-2-7-claims.png"),
        ("/ambulance-requests", "4-11-2-8-ambulance-req.png"),
        ("/admin/blogs", "4-11-2-9-blogs-admin.png"),
        ("/site-settings", "4-11-2-10-settings.png"),
        ("/queue", "4-11-2-11-queue.png"),
        ("/doctor-schedules", "4-11-2-12-schedules.png"),
    ]:
        page.goto(f"{BASE}{path}", wait_until="domcontentloaded")
        clean_ui(page)
        page.wait_for_timeout(500)
        shot(page, name)

    try:
        page.goto(f"{BASE}/dashboard/admin", wait_until="domcontentloaded")
        clean_ui(page)
        bell = page.locator(".dropdown-toggle, .nav-link, a").filter(has_text="").locator("i.fa-bell, .fe-bell, .typcn-bell").first
        if bell.count():
            bell.click(timeout=2000)
            page.wait_for_timeout(400)
        shot(page, "4-11-8-2-notifications.png")
    except Exception:
        shot(page, "4-11-8-2-notifications.png")

    # Doctor
    login(page, "doctor")
    for path, name in [
        ("/doctor/appointments", "4-11-3-1-doctor-appointments.png"),
        ("/doctor/queue", "4-11-3-2-doctor-queue.png"),
        ("/doctor/invoices", "4-11-3-3-diagnosis.png"),
        ("/doctor/invoices", "4-11-3-4-labs-rays.png"),
        ("/doctor/patient_details/10", "4-11-3-5-patient-details.png"),
        ("/doctor/chat/patients", "4-11-3-6-chat.png"),
    ]:
        page.goto(f"{BASE}{path}", wait_until="domcontentloaded")
        clean_ui(page)
        page.wait_for_timeout(600)
        if name == "4-11-3-3-diagnosis.png":
            try:
                page.evaluate("document.querySelector('[data-target^=\"#add_diagnosis\"]')?.click()")
                page.wait_for_timeout(600)
            except Exception as exc:
                print("diagnosis modal", exc)
        if name == "4-11-3-4-labs-rays.png":
            try:
                page.evaluate("document.querySelector('[data-target^=\"#xray_conversion\"]')?.click()")
                page.wait_for_timeout(600)
            except Exception as exc:
                print("ray modal", exc)
        shot(page, name)

    # Patient
    login(page, "patient")
    page.goto(f"{BASE}/my-appointments", wait_until="domcontentloaded")
    clean_ui(page)
    shot(page, "4-11-4-1-my-appointments.png")
    try:
        rate = page.locator("a", has_text="قيّم الطبيب").first
        if rate.count():
            rate.click()
            page.wait_for_load_state("domcontentloaded")
            page.wait_for_timeout(500)
        else:
            page.goto(f"{BASE}/my-appointments", wait_until="domcontentloaded")
        clean_ui(page)
        shot(page, "4-11-4-2-rating.png")
    except Exception as exc:
        print("rating page", exc)
        shot(page, "4-11-4-2-rating.png")
    page.goto(f"{BASE}/rays", wait_until="domcontentloaded")
    clean_ui(page)
    shot(page, "4-11-4-3-results.png")
    page.goto(f"{BASE}/laboratories", wait_until="domcontentloaded")
    clean_ui(page)
    shot(page, "4-11-4-4-pdf.png")
    page.goto(f"{BASE}/invoices", wait_until="domcontentloaded")
    clean_ui(page)
    shot(page, "4-11-4-5-invoices.png")

    # Ray employee
    login(page, "ray_employee")
    page.goto(f"{BASE}/invoices_ray_employee", wait_until="domcontentloaded")
    clean_ui(page)
    shot(page, "4-11-5-1-ray-list.png")
    page.goto(f"{BASE}/invoices_ray_employee/7/edit", wait_until="domcontentloaded")
    clean_ui(page)
    shot(page, "4-11-5-2-ray-edit.png")

    # Lab employee
    login(page, "laboratorie_employee")
    page.goto(f"{BASE}/invoices_laboratorie_employee", wait_until="domcontentloaded")
    clean_ui(page)
    shot(page, "4-11-6-1-lab-list.png")
    page.goto(f"{BASE}/invoices_laboratorie_employee/7/edit", wait_until="domcontentloaded")
    clean_ui(page)
    shot(page, "4-11-6-2-lab-edit.png")


def insert_into_docx() -> None:
    if not DOCX.exists():
        raise SystemExit(f"missing {DOCX}")
    shutil.copy2(DOCX, BACKUP)
    doc = Document(str(DOCX))
    pending = None
    inserted = 0
    unmatched = []
    for p in doc.paragraphs:
        text = " ".join((p.text or "").split())
        for heading, filename in HEADING_IMAGES.items():
            key = " ".join(heading.split())
            if text == key or text.startswith(key):
                pending = filename
                break
        if pending and ("أدخل لقطة" in text or "لقطة الشاشة" in text):
            img = SHOTS / pending
            if img.exists():
                p.clear()
                run = p.add_run()
                run.add_picture(str(img), width=Cm(14.8))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                inserted += 1
            else:
                unmatched.append(pending)
            pending = None
    alt = DOCX
    try:
        doc.save(str(DOCX))
    except PermissionError:
        alt = DOCS / "تطوير نظام ادارة المستشفى - لقطات.docx"
        doc.save(str(alt))
        print("original locked, saved", alt.name)
    print(f"inserted {inserted} images into {alt.name}")
    print(f"backup: {BACKUP.name}")
    if unmatched:
        print("missing images:", unmatched)


def main() -> None:
    SHOTS.mkdir(parents=True, exist_ok=True)
    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        context = browser.new_context(viewport={"width": 1440, "height": 900}, locale="ar")
        page = context.new_page()
        page.set_default_timeout(15000)
        try:
            capture(page)
        except Exception as exc:
            print("capture error:", exc)
        finally:
            browser.close()
    insert_into_docx()


if __name__ == "__main__":
    main()
