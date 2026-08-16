# -*- coding: utf-8 -*-
"""Third pass: TOC cleanup, section order, chapter 5 header."""
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = 'تطوير نظام ادارة المستشفى - كامل.docx'

SECTIONS = [
    ('4.7.6 الموقع الإلكتروني العام', 'Side title', True),
    ('صفحة رئيسية (Hero، إحصائيات، أقسام، أطباء، مدونة)، حجز مواعيد AJAX، طلب إسعاف، إعدادات الموقع.', 'arabic', False),
    ('4.7.7 نظام الإشعارات والتواصل', 'Side title', True),
    ('NotificationService، بريد تأكيد الموعد، SMS Twilio (اختياري).', 'arabic', False),
    ('4.7.8 الوصفات الإلكترونية (e-Prescription)', 'Side title', True),
    ('جدول prescriptions: دواء، جرعة، تكرار، مدة، تعليمات — مرتبط بالتشخيص.', 'arabic', False),
    ('4.7.9 لوحة التقارير والإحصائيات', 'Side title', True),
    ('مسار /reports — Chart.js: مرضى/شهر، إيرادات، أداء أقسام.', 'arabic', False),
    ('4.7.10 جدولة أوقات الطبيب', 'Side title', True),
    ('doctor_schedules + AppointmentScheduleService لمنع التعارض.', 'arabic', False),
    ('4.7.11 تكامل التأمين الكامل', 'Side title', True),
    ('insurance_id، خصم تلقائي، مطالبات insurance_claims، تقارير الشركات.', 'arabic', False),
    ('4.7.12 تذكير المواعيد', 'Side title', True),
    ('appointments:send-reminders — Email + SMS قبل 24 ساعة.', 'arabic', False),
    ('4.7.13 تقييم الأطباء', 'Side title', True),
    ('doctor_ratings: نجوم 1–5 بعد موعد منتهي.', 'arabic', False),
    ('4.7.14 إدارة المدونة', 'Side title', True),
    ('CRUD /admin/blogs — إعجاب/تعليق للمريض، مودال للزائر.', 'arabic', False),
    ('4.7.15 طلب الإسعاف', 'Side title', True),
    ('ambulance_requests — نموذج الموقع → إرسال سيارة.', 'arabic', False),
    ('4.7.16 تصدير السجل الطبي PDF', 'Side title', True),
    ('DomPDF — تشخيصات + وصفات + أشعة + مختبر.', 'arabic', False),
    ('4.7.17 إدارة الطابور (Queue Management)', 'Side title', True),
    ('queue_tickets — /queue، /doctor/queue، شاشة TV، /queue/track، Pusher.', 'arabic', False),
    ('4.7.18 بيانات تجريبية (Database Seeders)', 'Side title', True),
    ('HmsFullDemoSeeder — php artisan db:seed | migrate:fresh --seed', 'arabic', False),
    ('admin@gmail.com/123456789 — patient@yahoo.com/12345678', 'arabic', False),
]

TOC_47 = [s[0] for s in SECTIONS if s[1] == 'Side title']
TOC_CH5 = [
    ('الفصل الخامس — الاختبار', 'toc 1'),
    ('5.1 أنواع الاختبارات', 'toc 3'),
    ('5.2 سيناريوهات الاختبار', 'toc 3'),
    ('5.2.5 اختبار الميزات المتقدمة', 'toc 3'),
    ('5.3 نتائج الاختبار', 'toc 3'),
    ('5.4 خاتمة الفصل', 'toc 3'),
]


def delete_paragraph(para):
    para._element.getparent().remove(para._element)


def body_start(doc):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'ملخص المشروع' and p.style.name == 'main title':
            return i
    return len(doc.paragraphs)


def insert_after(paragraph, text, style='arabic', bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def insert_before(paragraph, text, style='arabic', bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def clean_toc(doc):
    toc_i = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == 'فهرس المحتويات')
    body_i = body_start(doc)

    junk_patterns = (
        '- Models', '- Repository', '- Helper', '- تكامل', 'Results and Conclusions',
    )
    for para in list(doc.paragraphs)[toc_i + 1:body_i]:
        t = para.text.strip()
        if not para.style.name.startswith('toc'):
            continue
        if any(t.startswith(j) or t == j for j in junk_patterns):
            delete_paragraph(para)
        if t == 'الخلاصة والاستنتاجات':
            delete_paragraph(para)

    # Reorder 4.7.x ascending in TOC
    anchor = None
    for para in doc.paragraphs[toc_i:body_i]:
        if para.text.strip().startswith('4.8'):
            anchor = para
            break
    if anchor:
        for para in list(doc.paragraphs)[toc_i:body_i]:
            t = para.text.strip()
            if para.style.name.startswith('toc') and re.match(r'^4\.7\.\d+', t):
                delete_paragraph(para)
        prev = anchor
        for title in reversed(TOC_47):
            prev = insert_before(anchor, title, style='toc 3')

    # Add chapter 5 TOC if missing
    existing = {p.text.strip() for p in doc.paragraphs[toc_i:body_i]}
    anchor = None
    for para in doc.paragraphs[toc_i:body_i]:
        if para.text.strip().startswith('4.10'):
            anchor = para
            break
    if anchor:
        for title, style in reversed(TOC_CH5):
            if title not in existing:
                insert_after(anchor, title, style=style)

    # Add chapters 7-9 if missing
    for title, style in [
        ('الفصل الثامن — الأفاق المستقبلية', 'toc 1'),
        ('الفصل التاسع — المراجع', 'toc 1'),
    ]:
        if title not in existing:
            last = doc.paragraphs[body_i - 1]
            insert_before(last if last.style.name.startswith('toc') else doc.paragraphs[body_i], title, style=style)


def reorder_body_47(doc):
    body_i = body_start(doc)
    p75 = None
    p48 = None
    for para in doc.paragraphs[body_i:]:
        t = para.text.strip()
        if t.startswith('4.7.5'):
            p75 = para
        if t.startswith('4.8') and para.style.name in ('Side title', 'Sub-Side title'):
            p48 = para
            break
    if not p75 or not p48:
        return

    # delete existing 4.7.6-18 block
    deleting = False
    for para in list(doc.paragraphs)[body_i:]:
        t = para.text.strip()
        if t.startswith('4.7.5'):
            deleting = True
            continue
        if deleting:
            if t.startswith('4.8'):
                break
            if t or para.runs:
                delete_paragraph(para)

    p75 = None
    for para in doc.paragraphs[body_i:]:
        if para.text.strip().startswith('4.7.5'):
            p75 = para
            break
    if not p75:
        return

    last = p75
    for text, style, bold in SECTIONS:
        last = insert_after(last, text, style=style, bold=bold)


def ensure_chapter5_header(doc):
    body_i = body_start(doc)
    has_ch5 = any(
        'الفصل الخامس' in p.text and p.style.name == 'main title'
        for p in doc.paragraphs[body_i:]
    )
    if has_ch5:
        return
    anchor = None
    for para in doc.paragraphs[body_i:]:
        if para.text.strip().startswith('5.1'):
            anchor = para
            break
    if anchor:
        p = insert_before(anchor, 'الفصل الخامس — الاختبار', style='main title', bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        insert_before(anchor, 'Testing', style='English')


def main():
    doc = Document(SOURCE)
    clean_toc(doc)
    reorder_body_47(doc)
    ensure_chapter5_header(doc)
    doc.save(SOURCE)
    print('Pass 3 done.')


if __name__ == '__main__':
    main()
