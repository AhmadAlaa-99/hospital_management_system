# -*- coding: utf-8 -*-
"""Fix misplaced 12.7.6-12.7.16 blocks — move to after 12.7.5 in body."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_before(paragraph, text, bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def main():
    doc = Document('تطوير نظام ادارة المستشفى.docx')
    paras = list(doc.paragraphs)

    # Collect wrongly placed 12.7.6-12.7.16 from TOC (before first '12.7.1')
    toc_blocks = []
    in_block = False
    to_delete = []
    found_body = False
    for para in paras:
        t = para.text.strip()
        if t.startswith('12.7.1 '):
            found_body = True
        if not found_body and (t.startswith('12.7.6') or t.startswith('12.7.7') or t.startswith('12.7.8')
                              or t.startswith('12.7.9') or t.startswith('12.7.1') and '0 ' in t[:8]
                              or t.startswith('12.7.11') or t.startswith('12.7.12') or t.startswith('12.7.13')
                              or t.startswith('12.7.14') or t.startswith('12.7.15') or t.startswith('12.7.16')):
            # only 12.7.6+ in toc area
            pass

    # Simpler: delete any 12.7.6-16 before paragraph containing '12.7.1 إدارة'
    body_271_idx = None
    for i, para in enumerate(paras):
        if para.text.strip().startswith('12.7.1 ') and 'إدارة' in para.text:
            body_271_idx = i
            break

    if body_271_idx is None:
        print('Could not find body 12.7.1')
        return

    # Delete 12.7.6-16 paragraphs that appear BEFORE body_271_idx
    section_prefixes = tuple(f'12.7.{n}' for n in range(6, 17))
    for para in paras[:body_271_idx]:
        t = para.text.strip()
        if any(t.startswith(p) for p in section_prefixes):
            delete_paragraph(para)
        elif body_271_idx and any(x in t for x in [
            'NotificationService', 'DomPDF', 'doctor_schedules', 'insurance_claims',
            'appointments:send-reminders', 'doctor_ratings', '/admin/blogs',
            'ambulance_requests', 'صفحة رئيسية', 'جدول prescriptions',
        ]):
            delete_paragraph(para)

    # Re-scan paragraphs
    paras = list(doc.paragraphs)
    anchor_128 = None
    after_275 = False
    for para in paras:
        t = para.text.strip()
        if t.startswith('12.7.5'):
            after_275 = True
            continue
        if after_275 and t.startswith('12.8'):
            anchor_128 = para
            break

    if not anchor_128:
        print('Could not find 12.8 after 12.7.5')
        return

    # Check if 12.7.6 already exists between 12.7.5 and 12.8
    has_276 = False
    collecting = False
    for para in paras:
        t = para.text.strip()
        if t.startswith('12.7.5'):
            collecting = True
            continue
        if collecting and t.startswith('12.7.6'):
            has_276 = True
            break
        if collecting and t.startswith('12.8'):
            break

    if not has_276:
        blocks = [
            ('12.7.6 الموقع الإلكتروني العام', True),
            ('صفحة رئيسية، حجز مواعيد، مدونة، طلب إسعاف، إعدادات الموقع.', False),
            ('12.7.7 نظام الإشعارات والتواصل', True),
            ('NotificationService، بريد تأكيد، SMS Twilio، إشعارات فورية.', False),
            ('12.7.8 الوصفات الإلكترونية (e-Prescription)', True),
            ('جدول prescriptions: دواء، جرعة، تكرار، مدة، تعليمات — مرتبط بالتشخيص.', False),
            ('12.7.9 لوحة التقارير والإحصائيات', True),
            ('/reports — Charts: مرضى/شهر، إيرادات، أداء أقسام، متوسط تقييم الأطباء.', False),
            ('12.7.10 جدولة أوقات الطبيب', True),
            ('doctor_schedules + API slots + منع التعارض.', False),
            ('12.7.11 تكامل التأمين الكامل', True),
            ('insurance_id، خصم تلقائي، مطالبات، تقارير الشركات.', False),
            ('12.7.12 تذكير المواعيد', True),
            ('appointments:send-reminders — Email + SMS قبل 24 ساعة.', False),
            ('12.7.13 تقييم الأطباء', True),
            ('doctor_ratings: نجوم 1–5 وتعليق بعد موعد منتهي.', False),
            ('12.7.14 إدارة المدونة', True),
            ('CRUD من Dashboard: /admin/blogs.', False),
            ('12.7.15 طلب الإسعاف', True),
            ('نموذج الموقع → ambulance_requests → إرسال سيارة.', False),
            ('12.7.16 تصدير السجل الطبي PDF', True),
            ('DomPDF — تشخيصات + وصفات + أشعة + مختبر — للطبيب والمريض والمدير.', False),
        ]
        for text, bold in reversed(blocks):
            insert_before(anchor_128, text, bold=bold)

    # Fix remaining Gates in theory chapter
    for para in doc.paragraphs:
        if 'Gates' in para.text and 'Policies' in para.text:
            para.text = '- Middleware متعدد الحراس (Multi-guard) للتحكم بالصلاحيات'

    doc.save('تطوير نظام ادارة المستشفى.docx')
    doc.save('docs/تطوير نظام ادارة المستشفى - كامل.docx')
    print('Fixed placement')


if __name__ == '__main__':
    main()
