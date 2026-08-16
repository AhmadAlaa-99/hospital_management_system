# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document('تطوير نظام ادارة المستشفى - كامل.docx')

for i in range(55, 115):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if t:
        print(f'{i:4d}| {p.style.name:15s}| {t[:80]}')

print('--- body 870-930 ---')
for i in range(870, 930):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if t:
        print(f'{i:4d}| {p.style.name:15s}| {t[:80]}')

print('--- results list ---')
for i in range(1085, 1110):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if t:
        print(f'{i:4d}| {p.style.name:15s}| {t[:80]}')
