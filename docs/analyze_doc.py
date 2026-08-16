# -*- coding: utf-8 -*-
"""Analyze document structure for formatting fixes."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from collections import Counter

doc = Document('تطوير نظام ادارة المستشفى - كامل.docx')

print('=== Chapter title misuse ===')
count = 0
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Chapter title' and len(p.text.strip()) > 100:
        count += 1
        if count <= 10:
            print(f'{i:4d}| len={len(p.text)}| {p.text[:90]}...')
print(f'Total: {count}')

print('\n=== Cover page ===')
for i in range(0, 14):
    p = doc.paragraphs[i]
    print(f'{i:3d}| {p.style.name:15s}| {p.text[:60]!r}')

print('\n=== Key transition points ===')
keywords = ['12.10', '11.4', '11.5', '12. الفصل', '13. الفصل', '15.', '16.', '17.', 'الخلاصة']
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if any(t.startswith(k) or t == k for k in keywords):
        print(f'{i:4d}| {p.style.name:20s}| {t[:75]}')

print('\n=== Styles defined ===')
for s in doc.styles:
    if s.type is not None and s.name in ('arabic', 'Chapter title', 'main title', 'Side title', 'Sub-Side title', 'Normal', 'toc 1', 'toc 2', 'toc 3'):
        try:
            base = s.base_style.name if s.base_style else None
            print(f'  {s.name}: base={base}')
        except Exception:
            print(f'  {s.name}')
