# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from collections import Counter

doc = Document('تطوير نظام ادارة المستشفى - كامل.docx')

print('Paragraphs:', len(doc.paragraphs))

issues = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Tailwind' in t:
        issues.append((i, 'TAILWIND', t[:80]))
    if t in ('73', '79'):
        issues.append((i, 'ORPHAN', t))
    if t.startswith('11.4'):
        issues.append((i, 'DUP_11.4', t[:80]))
    if '12.7.' in t or t.startswith('12.'):
        issues.append((i, 'OLD_12', t[:80]))
    if t.startswith('13.'):
        issues.append((i, 'OLD_13', t[:80]))

print('\n=== Remaining issues ===')
for x in issues[:30]:
    print(x)
print(f'Total issues: {len(issues)}')

print('\n=== Cover ===')
for i in range(0, 16):
    p = doc.paragraphs[i]
    if p.text.strip():
        print(f'{i:3d}| {p.text.strip()}')

print('\n=== TOC (34-95) ===')
for i in range(34, min(95, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    if p.text.strip() and p.style.name.startswith('toc'):
        print(f'{i:3d}| {p.style.name:8s}| {p.text.strip()[:70]}')

print('\n=== 4.7 sections in body ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('4.7.'):
        print(f'{i:4d}| {p.style.name:15s}| {t[:75]}')

print('\n=== Chapter headers ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'الفصل' in t and p.style.name == 'main title':
        print(f'{i:4d}| {t}')

print('\n=== Summary style check ===')
idx = next(i for i,p in enumerate(doc.paragraphs) if p.text.strip()=='ملخص المشروع')
for i in range(idx+1, idx+8):
    p = doc.paragraphs[i]
    if p.text.strip():
        print(f'{i}| {p.style.name}| {p.text[:60]}...')
