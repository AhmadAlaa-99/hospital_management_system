# -*- coding: utf-8 -*-
"""Add Queue Management (12.7.17) to graduation Word document."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_before(paragraph, text, bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def insert_after(paragraph, text, bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def find_paragraph(doc, startswith=None, equals=None, contains=None):
    for para in doc.paragraphs:
        t = para.text.strip()
        if equals and t == equals:
            return para
        if startswith and t.startswith(startswith):
            return para
        if contains and contains in t:
            return para
    return None


def already_has_queue(doc):
    return find_paragraph(doc, startswith='12.7.17') is not None


def main():
    paths = [
        'تطوير نظام ادارة المستشفى.docx',
        'docs/تطوير نظام ادارة المستشفى - كامل.docx',
    ]

    for path in paths:
        try:
            doc = Document(path)
        except Exception as e:
            print(f'Skip {path}: {e}')
            continue

        if already_has_queue(doc):
            print(f'Already updated: {path}')
            doc.save(path)
            continue

        anchor = find_paragraph(doc, startswith='12.8')
        if anchor:
            blocks = [
                ('12.7.17 إدارة الطابور (Queue Management)', True),
                ('جدول queue_tickets — إصدار أرقام، تسجيل حضور من الموعد، نداء التالي،', False),
                ('شاشة عرض TV، تتبع عام للمريض، بث Pusher + polling.', False),
                ('المسارات: /queue (مدير)، /doctor/queue (طبيب)، /queue/track (عام).', False),
            ]
            for text, bold in reversed(blocks):
                insert_before(anchor, text, bold=bold)

        p16 = find_paragraph(doc, contains='16.0 الميزات المنفّذة')
        if p16:
            insert_after(p16, 'إدارة الطابور (Queue Management) مع شاشة عرض وتتبع.', False)

        p131 = find_paragraph(doc, contains='14. تصدير السجل الطبي PDF')
        if p131:
            insert_after(p131, '15. إدارة الطابور: استقبال، طبيب، شاشة عرض، تتبع', False)

        for para in doc.paragraphs:
            if para.text.strip() == '5 | Queue Management | مقترح':
                para.text = '5 | Queue Management | منفّذ (12.7.17)'
            if 'Queue Management' in para.text and 'مقترح' in para.text and '|' in para.text:
                para.text = para.text.replace('مقترح', 'منفّذ')

        doc.save(path)
        print(f'Updated: {path}')


if __name__ == '__main__':
    main()
