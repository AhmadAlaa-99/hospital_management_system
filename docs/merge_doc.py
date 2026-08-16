# -*- coding: utf-8 -*-
"""Merge new features into original Word document sections."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_before(paragraph, text, bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def insert_after(paragraph, text, bold=False):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def find_paragraph(doc, startswith=None, equals=None, contains=None):
    for para in doc.paragraphs:
        t = para.text.strip()
        if equals and t == equals:
            return para
        if startswith and t.startswith(startswith):
            return para
        if contains and contains in t:
            return para
    return None


def main():
    path = 'تطوير نظام ادارة المستشفى.docx'
    doc = Document(path)

    # 1. Global replacements
    for para in doc.paragraphs:
        text = para.text
        text = text.replace('Livewire و Tailwind CSS', 'Livewire و Bootstrap (قالب Valex)')
        text = text.replace('Tailwind CSS', 'Bootstrap (قالب Valex)')
        text = text.replace(
            'Blade Templates و Livewire و Tailwind CSS',
            'Blade Templates و Livewire و Bootstrap (قالب Valex)',
        )
        if para.text != text:
            para.text = text

    # 2. Fix authorization — remove Gates/Policies bullets
    for para in list(doc.paragraphs):
        t = para.text.strip()
        if t in ('Gates و Policies:', '- Gates للتحقق من الصلاحيات العامة',
                 '- Policies للتحكم في الوصول للموارد المحددة'):
            delete_paragraph(para)

    auth = find_paragraph(doc, equals='- `auth.patient`: التحقق من تسجيل دخول المريض')
    if auth:
        insert_after(auth, '- Middleware متعدد الحراس (Multi-guard) يفصل صلاحيات كل دور')

    # 3. Remove appendix at end
    start_del = False
    for para in list(doc.paragraphs):
        if para.text.strip().startswith('الفصل 12.7 — الوحدات المتقدمة'):
            start_del = True
        if start_del:
            delete_paragraph(para)

    # 4. Enhance 12.7.1
    p = find_paragraph(doc, equals='- تعيين المواعيد المتاحة')
    if p:
        insert_after(p, '- جدول عمل الأطباء (doctor_schedules): أيام وساعات ومدة الموعد')

    # 5. Enhance 12.7.2
    p = find_paragraph(doc, equals='- إرسال إشعارات للمرضى والأطباء')
    if p:
        last = p
        for line in [
            '- حجز مواعيد من الموقع مع اختيار التاريخ والوقت المتاح',
            '- منع تعارض المواعيد (AppointmentScheduleService)',
            '- تذكير تلقائي قبل 24 ساعة (Email + SMS)',
            '- صفحة المواعيد المنتهية والمرفوضة',
        ]:
            last = insert_after(last, line)

    # 6. Enhance 12.7.3
    p = find_paragraph(doc, equals='- تقارير مالية مفصلة')
    if p:
        p.text = '- لوحة تقارير وإحصائيات (/reports) مع Chart.js'
        last = p
        for line in [
            '- مطالبات التأمين (insurance_claims) تُنشأ تلقائياً',
            '- تقرير مطالبات التأمين حسب شركة التأمين',
        ]:
            last = insert_after(last, line)

    # 7. Insert 12.7.6–12.7.16 before 12.8
    anchor = find_paragraph(doc, startswith='12.8')
    if anchor:
        blocks = [
            ('12.7.6 الموقع الإلكتروني العام', True),
            ('صفحة رئيسية، حجز مواعيد، مدونة، طلب إسعاف، إعدادات الموقع.', False),
            ('12.7.7 نظام الإشعارات والتواصل', True),
            ('NotificationService، بريد تأكيد، SMS Twilio، إشعارات فورية.', False),
            ('12.7.8 الوصفات الإلكترونية (e-Prescription)', True),
            ('جدول prescriptions: دواء، جرعة، تكرار، مدة، تعليمات — مرتبط بالتشخيص.', False),
            ('12.7.9 لوحة التقارير والإحصائيات', True),
            ('/reports — Charts: مرضى/شهر، إيرادات، أداء أقسام، متوسط تقييم الأطباء.', False),
            ('12.7.10 جدولة أوقات الطبيب', True),
            ('doctor_schedules + API slots + منع التعارض.', False),
            ('12.7.11 تكامل التأمين الكامل', True),
            ('insurance_id، خصم تلقائي، مطالبات، تقارير الشركات.', False),
            ('12.7.12 تذكير المواعيد', True),
            ('appointments:send-reminders — Email + SMS قبل 24 ساعة.', False),
            ('12.7.13 تقييم الأطباء', True),
            ('doctor_ratings: نجوم 1–5 وتعليق بعد موعد منتهي.', False),
            ('12.7.14 إدارة المدونة', True),
            ('CRUD من Dashboard: /admin/blogs.', False),
            ('12.7.15 طلب الإسعاف', True),
            ('نموذج الموقع → ambulance_requests → إرسال سيارة.', False),
            ('12.7.16 تصدير السجل الطبي PDF', True),
            ('DomPDF — تشخيصات + وصفات + أشعة + مختبر.', False),
            ('12.7.17 إدارة الطابور (Queue Management)', True),
            ('queue_tickets — استقبال، طبيب، شاشة TV، تتبع عام، Pusher.', False),
        ]
        for text, bold in reversed(blocks):
            insert_before(anchor, text, bold=bold)

    # 8. Update 13.1.1
    p = find_paragraph(doc, equals='- نظام إشعارات شامل')
    if p:
        last = p
        extras = [
            '6. الموقع الإلكتروني: حجز، مدونة، إسعاف، إعدادات',
            '7. الوصفات الإلكترونية e-Prescription',
            '8. لوحة التقارير والإحصائيات (Chart.js)',
            '9. جدولة الأطباء ومنع تعارض المواعيد',
            '10. تكامل التأمين (مطالبات + تقارير)',
            '11. تذكير المواعيد Email/SMS',
            '12. تقييم الأطباء',
            '13. إدارة المدونة CRUD',
            '14. تصدير السجل الطبي PDF',
        ]
        for line in extras:
            last = insert_after(last, line)

    # 9. Update 13.1.2
    for para in doc.paragraphs:
        if 'نظام أمان متقد' in para.text:
            last = para
            for line in ['-  DomPDF لتصدير السجلات', '-  Chart.js للتقارير', '-  Laravel Scheduler للتذكير']:
                last = insert_after(last, line)
            break

    # 10. Update 13.4.1 item 2
    p = find_paragraph(doc, contains='2. تحسين التقارير')
    if p:
        p.text = '2. توسيع التقارير: Excel/PDF (تم تنفيذ لوحة Chart.js الأساسية)'

    # 11. Section 16.0 before 16.1
    anchor16 = find_paragraph(doc, startswith='16.1')
    if anchor16:
        blocks16 = [
            ('16.0 الميزات المنفّذة', True),
            ('لوحة تقارير، e-Prescription، جدولة أطباء، تأمين كامل، تذكير مواعيد،', False),
            ('تقييم أطباء، مدونة CRUD، طلب إسعاف، تصدير PDF.', False),
        ]
        for text, bold in reversed(blocks16):
            insert_before(anchor16, text, bold=bold)

    doc.save('تطوير نظام ادارة المستشفى.docx')
    doc.save('docs/تطوير نظام ادارة المستشفى - كامل.docx')
    print('Done')


if __name__ == '__main__':
    main()
