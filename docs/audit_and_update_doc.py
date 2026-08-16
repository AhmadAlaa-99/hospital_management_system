# -*- coding: utf-8 -*-
"""
تدقيق وتحديث وثيقة مشروع التخرج (body فقط — لا يُعدّل الفهرس).
"""
from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOCS = Path(__file__).resolve().parent


def resolve_doc_path() -> Path:
    matches = list(DOCS.glob('*كامل*.docx'))
    if not matches:
        raise FileNotFoundError('Word document not found in docs/')
    return matches[0]


def is_toc(para: Paragraph) -> bool:
    return para.style.name.lower().startswith('toc')


def find_body_paragraph(doc: Document, **kwargs) -> Paragraph | None:
    for para in doc.paragraphs:
        if is_toc(para):
            continue
        t = para.text.strip()
        if kwargs.get('equals') is not None and t == kwargs['equals']:
            return para
        if kwargs.get('startswith') and t.startswith(kwargs['startswith']):
            return para
        if kwargs.get('contains') and kwargs['contains'] in t:
            return para
        if kwargs.get('style') and para.style.name == kwargs['style'] and kwargs.get('startswith'):
            if t.startswith(kwargs['startswith']):
                return para
    return None


def insert_before(paragraph: Paragraph, text: str, bold: bool = False, style: str | None = None) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def insert_after(paragraph: Paragraph, text: str, bold: bool = False, style: str | None = None) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    run = new_para.add_run(text)
    run.bold = bold
    return new_para


def replace_paragraph_text(para: Paragraph, new_text: str, bold: bool | None = None) -> None:
    b = bold if bold is not None else any(r.bold for r in para.runs)
    para.clear()
    run = para.add_run(new_text)
    run.bold = b


def replace_all_text(doc: Document, mapping: dict[str, str], body_only: bool = True) -> int:
    count = 0
    for para in doc.paragraphs:
        if body_only and is_toc(para):
            # update TOC titles for terminology only
            pass
        old = para.text
        new = old
        for k, v in mapping.items():
            new = new.replace(k, v)
        if new != old:
            replace_paragraph_text(para, new)
            count += 1
    return count


def delete_paragraph(para: Paragraph) -> None:
    para._element.getparent().remove(para._element)


def insert_blocks_before(anchor: Paragraph, blocks: list[tuple[str, bool, str | None]]) -> None:
    for text, bold, style in reversed(blocks):
        if not text:
            insert_before(anchor, ' ', bold=False)
            continue
        insert_before(anchor, text, bold=bold, style=style)


TERMINOLOGY_MAP = {
    '4.7.17 إدارة الطابور (Queue Management)': '4.7.17 نظام إدارة الانتظار (Queue Management)',
    'إدارة الطابور (Queue Management)': 'نظام إدارة الانتظار (Queue Management)',
    '15. إدارة الطابور: استقبال، طبيب، شاشة عرض، تتبع': '15. نظام إدارة الانتظار: استقبال، طبيب، شاشة عرض، تتبع',
    'إصدار رقم طابور': 'إصدار رقم انتظار (CARD-003)',
    'أرقام طابور': 'أرقام انتظار',
    'شاشة الطابور': 'شاشة الانتظار',
}

SECTION_47_EXPANSIONS: dict[str, list[str]] = {
    '4.7.6': [
        'يوفر الموقع العام واجهة تعريفية للمستشفى تشمل: صفحة رئيسية (Hero، إحصائيات، أقسام، أطباء، مقالات المدونة)، صفحة «من نحن»، صفحة الأقسام والأطباء، وحجز موعد عبر AJAX مع اختيار القسم والطبيب والتاريخ/الوقت المتاح.',
        'المسارات: / (رئيسية)، /appointment (حجز)، /blogs (المدونة)، /ambulance (طلب إسعاف)، /queue/track (تتبع الانتظار).',
        'تسجيل/دخول المريض من الموقع (WebsitePatientAuthController) لتمكين الحجز والتعليق والتقييم.',
        'إعدادات الموقع (SiteSetting): اسم المستشفى، الشعار، بيانات التواصل — قابلة للتعديل من لوحة المدير.',
    ],
    '4.7.7': [
        'NotificationService: إنشاء إشعارات داخلية حسب نوع المستخدم (admin، doctor، patient، ray_employee، laboratorie_employee).',
        'بريد تأكيد الموعد (AppointmentConfirmation) عند قبول الموعد من المدير.',
        'أمر appointments:send-reminders (Laravel Scheduler) يرسل Email وSMS (Twilio اختياري) قبل 24 ساعة.',
        'إشعارات تلقائية عند: نتائج أشعة/مختبر، موعد جديد، رفض موعد، طلب إسعاف.',
    ],
    '4.7.8': [
        'جدول prescriptions مرتبط بجدول diagnostics: drug_name، dosage، frequency، duration_days، instructions.',
        'يُضاف من لوحة الطبيب عند التشخيص (DiagnosisRepository) ويظهر في السجل الطبي وتصدير PDF.',
    ],
    '4.7.9': [
        'ReportsController — مسار /reports (مدير): رسوم Chart.js لمرضى جدد شهرياً، الإيرادات، أداء الأقسام.',
        'مؤشرات: إجمالي المرضى، الإيرادات، المواعيد، متوسط تقييم الأطباء، مطالبات التأمين حسب الشركة.',
    ],
    '4.7.10': [
        'جدول doctor_schedules: day_of_week، start_time، end_time، slot_duration — يُدار من DoctorScheduleController.',
        'AppointmentScheduleService: validateSlot يمنع الحجز المزدوج، يحدّ عدد المواعيد المعلقة، يتحقق من نشاط الطبيب.',
        'API أوقات متاحة للموقع العام عند اختيار الطبيب والتاريخ.',
    ],
    '4.7.11': [
        'ربط المريض بشركة تأمين (insurance_id) — خصم نسبة التغطية تلقائياً في الفواتير (Livewire Invoices).',
        'InsuranceClaimService ينشئ مطالبة insurance_claims تلقائياً مع الفاتورة.',
        'InsuranceClaimController: عرض، موافقة، رفض، تقارير حسب شركة التأمين.',
    ],
    '4.7.12': [
        'Console Command: appointments:send-reminders — مجدول كل ساعة في Kernel.php.',
        'يرسل للمواعيد المؤكدة قبل 24 ساعة عبر Mail وSMS (إن وُجد Twilio).',
    ],
    '4.7.13': [
        'جدول doctor_ratings: rating (1–5)، comment، patient_id، doctor_id، appointment_id (فريد).',
        'DoctorRatingController: المريض يقيّم بعد موعد منتهي مرة واحدة فقط.',
        'يظهر متوسط التقييم في لوحة التقارير وصفحة الأطباء.',
    ],
    '4.7.14': [
        'BlogController (Dashboard): CRUD كامل — title، slug، excerpt، content، image، published.',
        'BlogController (Website): عرض، إعجاب (BlogLike)، تعليق (BlogComment) للمريض؛ مودال تسجيل للزائر.',
    ],
    '4.7.15': [
        'AmbulanceRequestPublicController: نموذج طلب من الموقع (اسم، هاتف، موقع، حالة طوارئ).',
        'AmbulanceRequestController (مدير): قائمة الطلبات، assignAmbulance، complete، cancel.',
        'تحديث حالة سيارة الإسعاف (available/busy) تلقائياً.',
    ],
    '4.7.16': [
        'MedicalRecordPdfController + barryvdh/laravel-dompdf v2.',
        'يشمل: بيانات المريض، التشخيصات، الوصفات، نتائج الأشعة والمختبر.',
        'متاح للطبيب والمريض والمدير عبر مسارات medical-record/pdf/{patientId}.',
    ],
    '4.7.17': [
        'نظام إدارة الانتظار (Queue Management) — جدول queue_tickets: ticket_number (مثل CARD-003)، section_id، doctor_id، appointment_id، priority، status، queue_date.',
        'حالات التذكرة: waiting ثم called ثم serving ثم completed (أو no_show).',
        'لوحة الاستقبال /queue (QueueController): إصدار رقم walk-in، تسجيل حضور من موعد، فلترة بالقسم.',
        'لوحة الطبيب /doctor/queue: نداء التالي (callNext)، إعادة النداء، بدء الكشف، إنهاء، no-show.',
        'شاشة الانتظار TV /queue/display/section/{id}: عرض الرقم الحالي والتالي، تبديل الأقسام، تحديث Pusher + polling.',
        'تتبع عام /queue/track: المريض يدخل رقمه لمعرفة موقعه وتقدير وقت الانتظار (Throttling 30/min).',
        'QueueService: عمليات ذرية، unique (queue_date, appointment_id)، تقدير انتظار ديناميكي، ربط walk-in بالطبيب عند النداء.',
    ],
    '4.7.18': [
        'Seeder رئيسي HmsFullDemoSeeder يستدعي: UserTableSeeder، AdminTableSeeder، SectionTableSeeder، DoctorTableSeeder، PatientTableSeeder، Ray/Lab employees، Service/Group، AmbulanceInsuranceSeeder، AppointmentBookingSeeder، InvoiceDemoSeeder، ExtendedFeaturesSeeder، BlogSiteSettingSeeder، QueueTicketSeeder.',
        'الأوامر: php artisan db:seed | php artisan migrate:fresh --seed',
        'حسابات: admin@gmail.com/123456789 — patient@yahoo.com/12345678',
        'شاشة الانتظار: /queue/display/section/1 — تتبع: /queue/track',
    ],
}

SECTION_47_19 = [
    ('4.7.19 آليات الحماية والموثوقية', True, 'Side title'),
    ('AppointmentScheduleService: validateSlot مع استثناء appointment_id عند التعديل، حد أقصى للمواعيد المعلقة، رفض الطبيب غير النشط.', False, 'arabic'),
    ('QueueService: DB::transaction للحضور والنداء، منع تذكرة مزدوجة لنفس الموعد (unique index)، التحقق من تطابق قسم الطبيب.', False, 'arabic'),
    ('Throttling: حجز المواعيد 10/دقيقة، تتبع الانتظار 30/دقيقة.', False, 'arabic'),
    ('QueueUpdated event (ShouldBroadcastNow) + polling كل 5 ثوانٍ لضمان تحديث شاشة الانتظار.', False, 'arabic'),
    ('تسجيل أخطاء (Log) في QueueService وAppointmentScheduleService لتسهيل التشخيص.', False, 'arabic'),
]

SCREENSHOTS_SECTION = [
    ('4.11 لقطات الشاشة', True, 'Side title'),
    ('يوضح هذا القسم واجهات النظام المنفّذة. يُرجى إدراج لقطة الشاشة تحت كل عنوان.', False, 'arabic'),
    ('4.11.1 الموقع الإلكتروني العام', True, 'Sub-Side title'),
    ('4.11.1.1 الصفحة الرئيسية', False, 'arabic'),
    ('[أدخل لقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.1.2 صفحة حجز موعد', False, 'arabic'),
    ('[أدخل لقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.1.3 صفحة المدونة (قائمة المقالات)', False, 'arabic'),
    ('[أدخل لقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.1.4 صفحة تفاصيل مقال (إعجاب/تعليق)', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.1.5 صفحة طلب الإسعاف', False, 'arabic'),
    ('[أدخل لقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.1.6 صفحة تتبع الانتظار (/queue/track)', False, 'arabic'),
    ('[أدخل لقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.1.7 تسجيل/دخول المريض من الموقع', False, 'arabic'),
    ('[أدخل لقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2 لوحة تحكم المدير', True, 'Sub-Side title'),
    ('4.11.2.1 لوحة المعلومات الرئيسية (Dashboard)', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.2 إدارة الأقسام', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.3 إدارة الأطباء', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.4 إدارة المرضى', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.5 إدارة المواعيد', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.6 لوحة التقارير والإحصائيات (/reports)', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.7 إدارة مطالبات التأمين', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.8 إدارة طلبات الإسعاف', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.9 إدارة المدونة (CRUD)', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.10 إعدادات الموقع', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.11 إدارة الانتظار — لوحة الاستقبال (/queue)', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.2.12 جدولة أوقات الأطباء', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.3 لوحة تحكم الطبيب', True, 'Sub-Side title'),
    ('4.11.3.1 مواعيد الطبيب (قائمة/منتهية)', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.3.2 قائمة انتظار العيادة (/doctor/queue)', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.3.3 إضافة تشخiص ووصفة إلكترونية', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.3.4 طلب فحص أشعة/مختبر', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.3.5 تفاصيل المريض والسجل الطبي', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.3.6 المحادثة مع المريض (Livewire Chat)', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.4 لوحة المريض', True, 'Sub-Side title'),
    ('4.11.4.1 مواعيدي', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.4.2 تقييم الطبيب بعد الموعد', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.4.3 عرض نتائج الأشعة والمختبر', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.4.4 تصدير السجل الطبي PDF', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.4.5 الفواتير والمدفوعات', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.5 موظف الأشعة', True, 'Sub-Side title'),
    ('4.11.5.1 قائمة طلبات الأشعة', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.5.2 إدخال نتيجة الأشعة', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.6 موظف المختبر', True, 'Sub-Side title'),
    ('4.11.6.1 قائمة طلبات المختبر', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.6.2 إدخال نتيجة المختبر', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.7 شاشة الانتظار (TV)', True, 'Sub-Side title'),
    ('4.11.7.1 شاشة عرض الانتظار لقسم (/queue/display/section/{id})', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.7.2 تبديل الأقسام على شاشة العرض', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.8 المصادقة والأدوار', True, 'Sub-Side title'),
    ('4.11.8.1 صفحة تسجيل الدخول (لوحات التحكم)', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
    ('4.11.8.2 الإشعارات الداخلية', False, 'arabic'),
    ('[أدخل lقطة الشاشة هنا]', False, 'arabic'),
]

# fix typos in SCREENSHOTS_SECTION
SCREENSHOTS_SECTION = [
    (t.replace('lقطة', 'لقطة').replace('تشخiص', 'تشخيص'), b, s)
    for t, b, s in SCREENSHOTS_SECTION
]


def expand_section(doc: Document, section_key: str, lines: list[str]) -> None:
    title = find_body_paragraph(doc, startswith=section_key)
    if not title or title.style.name != 'Side title':
        # fallback: any body Side title matching
        for para in doc.paragraphs:
            if is_toc(para):
                continue
            if para.text.strip().startswith(section_key) and para.style.name == 'Side title':
                title = para
                break
    if not title:
        print(f'  skip expand {section_key}: not found')
        return

    to_delete: list[Paragraph] = []
    found = False
    for para in doc.paragraphs:
        if is_toc(para):
            continue
        t = para.text.strip()
        if para._p is title._p:
            found = True
            continue
        if found:
            if re.match(r'^4\.7\.\d+', t) and para.style.name == 'Side title':
                break
            if t.startswith('4.8') or t.startswith('Sub-Side title|4.8'):
                break
            if para.style.name == 'Sub-Side title' and t.startswith('4.8'):
                break
            if t.startswith('4.8'):
                break
            if t:
                to_delete.append(para)
    for p in to_delete:
        delete_paragraph(p)

    anchor = title
    for line in lines:
        anchor = insert_after(anchor, line, style='arabic')


def add_toc_entries(doc: Document) -> None:
    for para in doc.paragraphs:
        if not is_toc(para):
            continue
        t = para.text.strip()
        if '4.7.17' in t:
            replace_paragraph_text(para, t.replace('إدارة الطابور', 'نظام إدارة الانتظار'))
        if t.startswith('4.7.18') and not any(
            p.text.strip().startswith('4.7.19') for p in doc.paragraphs if is_toc(p)
        ):
            nxt = insert_after(para, '4.7.19 آليات الحماية والموثوقية')
            try:
                nxt.style = 'toc 3'
            except KeyError:
                pass
        if t.startswith('4.10 خاتمة') and not any(
            p.text.strip().startswith('4.11') for p in doc.paragraphs if is_toc(p)
        ):
            nxt = insert_after(para, '4.11 لقطات الشاشة')
            try:
                nxt.style = 'toc 3'
            except KeyError:
                pass


def update_testing_section(doc: Document) -> None:
    p = find_body_paragraph(doc, contains='queue/track')
    if p and 'طابور' in p.text:
        replace_paragraph_text(p, p.text.replace('طابور', 'انتظار (CARD-003)'))
    anchor = find_body_paragraph(doc, contains='تتبع الانتظار')
    if not anchor:
        anchor = find_body_paragraph(doc, contains='queue/track')
    extra_tests = [
        '- تسجيل حضور موعد في إدارة الانتظار (/queue/check-in)',
        '- نداء التالي من لوحة الطبيب (/doctor/queue/call-next)',
        '- عرض شاشة الانتظار TV مع تبديل الأقسام',
        '- التحقق من منع حجز موعد متعارض (validateSlot)',
        '- تشغيل HmsFullDemoSeeder واختبار الحسابات التجريبية',
    ]
    if anchor:
        last = anchor
        for line in extra_tests:
            if not find_body_paragraph(doc, contains=line[2:35]):
                last = insert_after(last, line, style='arabic')


def remove_orphan_lines_ch8(doc: Document) -> None:
    orphans = (
        'تقييم أطباء، مدونة CRUD مع إعجاب/تعليق، طلب إسعاف، تصدير PDF.',
        'لوحة تقارير، e-Prescription، جدولة أطباء، تأمين كامل، تذكير مواعيد،',
    )
    for para in doc.paragraphs:
        if is_toc(para):
            continue
        if para.text.strip() in orphans:
            delete_paragraph(para)


def process_document(path: Path) -> None:
    backup = path.with_name(f'backup_{datetime.now():%Y%m%d_%H%M%S}_audit.docx')
    shutil.copy2(path, backup)
    print(f'Backup: {backup.name}')

    doc = Document(str(path))

    t17 = find_body_paragraph(doc, startswith='4.7.17')
    if t17 and 'الطابور' in t17.text:
        replace_paragraph_text(t17, t17.text.replace('إدارة الطابور', 'نظام إدارة الانتظار'))

    n = replace_all_text(doc, TERMINOLOGY_MAP)
    print(f'Terminology replacements: {n}')

    for key, lines in SECTION_47_EXPANSIONS.items():
        expand_section(doc, key, lines)

    if not find_body_paragraph(doc, startswith='4.7.19'):
        anchor = find_body_paragraph(doc, startswith='4.8 الاختبارات')
        if anchor:
            insert_blocks_before(anchor, SECTION_47_19)
            print('Added 4.7.19')

    if not find_body_paragraph(doc, startswith='4.11 لقطات'):
        anchor = find_body_paragraph(doc, equals='الفصل الخامس — الاختبار')
        if not anchor:
            for para in doc.paragraphs:
                if is_toc(para):
                    continue
                if para.style.name == 'main title' and 'الفصل الخامس' in para.text:
                    anchor = para
                    break
        if anchor:
            insert_blocks_before(anchor, SCREENSHOTS_SECTION)
            print('Added 4.11 screenshots section')

    add_toc_entries(doc)
    update_testing_section(doc)
    remove_orphan_lines_ch8(doc)

    p80body = find_body_paragraph(doc, contains='لوحة تقارير، e-Prescription')
    if p80body:
        replace_paragraph_text(
            p80body,
            'لوحة تقارير، e-Prescription، جدولة أطباء، تأمين كامل، تذكير مواعيد، تقييم أطباء، '
            'مدونة CRUD، طلب إسعاف، تصدير PDF، نظام إدارة الانتظار (Queue Management)، '
            'Seeders تجريبية (HmsFullDemoSeeder)، آليات حماية الحجز والانتظار.',
        )

    p612 = find_body_paragraph(doc, contains='Laravel Scheduler للتذكير')
    extras = [
        '-  Pusher + QueueUpdated لشاشة الانتظار',
        '-  QueueService + AppointmentScheduleService لمنع التعارض',
    ]
    if p612:
        last = p612
        for line in extras:
            if not find_body_paragraph(doc, contains=line.strip()[2:28]):
                last = insert_after(last, line, style='arabic')

    doc.save(str(path))
    print(f'Saved: {path}')


def main() -> None:
    path = resolve_doc_path()
    process_document(path)
    for p in DOCS.parent.glob('تطوير*.docx'):
        if 'كامل' not in p.name:
            shutil.copy2(path, p)
            print(f'Copied to {p.name}')
            break


if __name__ == '__main__':
    main()
