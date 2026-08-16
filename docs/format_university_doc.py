# -*- coding: utf-8 -*-
"""
تنسيق وثيقة مشروع التخرج للتسليم الجامعي.
- إصلاح صفحة الغلاف والفهرس
- إزالة المحتوى المكرر والمُدرج في مكان خاطئ
- إعادة ترتيب 12.7.6–12.7.18
- توحيد أنماط الفقرات والترقيم
"""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

DOCS = Path(__file__).resolve().parent
SOURCE = DOCS / 'تطوير نظام ادارة المستشفى - كامل.docx'
OUTPUT = DOCS / 'تطوير نظام ادارة المستشفى - كامل.docx'
BACKUP = DOCS / f'backup_{datetime.now():%Y%m%d_%H%M%S}_تطوير نظام ادارة المستشفى.docx'

STUDENTS = ['صفا بديع الشهابي', 'ميسون حسين عبدالله']

# ── helpers ──────────────────────────────────────────────────────────

def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_before(paragraph: Paragraph, text: str, style: str = 'arabic', bold: bool = False) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    set_paragraph(new_para, text, style=style, bold=bold)
    return new_para


def insert_after(paragraph: Paragraph, text: str, style: str = 'arabic', bold: bool = False) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    set_paragraph(new_para, text, style=style, bold=bold)
    return new_para


def set_rtl(run) -> None:
    rpr = run._element.get_or_add_rPr()
    rtl = rpr.find(qn('w:rtl'))
    if rtl is None:
        rtl = OxmlElement('w:rtl')
        rpr.append(rtl)
    rtl.set(qn('w:val'), '1')


def set_paragraph(
    para: Paragraph,
    text: str,
    *,
    style: str = 'arabic',
    bold: bool = False,
    align=None,
    size_pt: float | None = None,
) -> None:
    try:
        para.text = ''
    except AttributeError:
        for run in list(para.runs):
            run.text = ''
    try:
        para.style = style
    except KeyError:
        para.style = 'Normal'
    if para.runs:
        para.runs[0].text = text
        run = para.runs[0]
        for r in para.runs[1:]:
            r.text = ''
    else:
        run = para.add_run(text)
    run.bold = bold
    set_rtl(run)
    if size_pt:
        run.font.size = Pt(size_pt)
        run.font.name = 'Traditional Arabic'
    if align is not None:
        para.alignment = align


def find_paragraph(doc: Document, *, startswith: str | None = None, equals: str | None = None, contains: str | None = None):
    for para in doc.paragraphs:
        t = para.text.strip()
        if equals is not None and t == equals:
            return para
        if startswith is not None and t.startswith(startswith):
            return para
        if contains is not None and contains in t:
            return para
    return None


def find_index(doc: Document, **kwargs) -> int | None:
    para = find_paragraph(doc, **kwargs)
    if para is None:
        return None
    for i, p in enumerate(doc.paragraphs):
        if p._p is para._p:
            return i
    return None


def strip_toc_page_num(text: str) -> str:
    return re.sub(r'\t+\d+\s*$', '', text.strip())


def replace_in_paragraph(para: Paragraph, mapping: dict[str, str]) -> None:
    text = para.text
    new = text
    for old, new_val in mapping.items():
        new = new.replace(old, new_val)
    if new != text:
        bold = any(r.bold for r in para.runs) if para.runs else False
        style = para.style.name
        set_paragraph(para, new, style=style, bold=bold)


def replace_all(doc: Document, mapping: dict[str, str]) -> None:
    for para in doc.paragraphs:
        replace_in_paragraph(para, mapping)


# ── content blocks ───────────────────────────────────────────────────

SECTIONS_127_ADVANCED = [
    ('4.7.6 الموقع الإلكتروني العام', 'Side title', True, None),
    ('صفحة رئيسية (Hero، إحصائيات، أقسام، أطباء، مدونة)، حجز مواعيد AJAX، طلب إسعاف، إعدادات الموقع.', 'arabic', False, None),
    ('4.7.7 نظام الإشعارات والتواصل', 'Side title', True, None),
    ('NotificationService للإشعارات الداخلية، بريد تأكيد الموعد (AppointmentConfirmation)، SMS عبر Twilio (اختياري)، تذكير قبل 24 ساعة.', 'arabic', False, None),
    ('4.7.8 الوصفات الإلكترونية (e-Prescription)', 'Side title', True, None),
    ('جدول prescriptions مرتبط بالتشخيص: اسم الدواء، الجرعة، التكرار، المدة (أيام)، التعليمات. يظهر في السجل الطبي وتصدير PDF.', 'arabic', False, None),
    ('4.7.9 لوحة التقارير والإحصائيات', 'Side title', True, None),
    ('مسار /reports للمدير — رسوم Chart.js: مرضى جدد شهرياً، إيرادات، أداء الأقسام. مؤشرات: إجمالي المرضى، الإيرادات، المواعيد، متوسط تقييم الأطباء.', 'arabic', False, None),
    ('4.7.10 جدولة أوقات الطبيب', 'Side title', True, None),
    ('جدول doctor_schedules (يوم، من، إلى، مدة الموعد). الموقع يعرض الأوقات المتاحة. AppointmentScheduleService يمنع تعارض المواعيد.', 'arabic', False, None),
    ('4.7.11 تكامل التأمين الكامل', 'Side title', True, None),
    ('ربط المريض بشركة تأمين (insurance_id)، خصم تلقائي في الفواتير، مطالبات insurance_claims تُنشأ تلقائياً، إدارة الحالات وتقرير حسب شركة التأمين.', 'arabic', False, None),
    ('4.7.12 تذكير المواعيد', 'Side title', True, None),
    ('أمر appointments:send-reminders مجدول كل ساعة عبر Laravel Scheduler — Email + SMS قبل 24 ساعة من الموعد المؤكد.', 'arabic', False, None),
    ('4.7.13 تقييم الأطباء', 'Side title', True, None),
    ('جدول doctor_ratings: تقييم 1–5 وتعليق. المريض يقيّم بعد موعد منتهي (مرة واحدة). يظهر المتوسط في لوحة التقارير.', 'arabic', False, None),
    ('4.7.14 إدارة المدونة', 'Side title', True, None),
    ('CRUD كامل من Dashboard (/admin/blogs): عنوان، slug، ملخص، محتوى، صورة، نشر/مسودة. إعجاب وتعليق للمريض المسجّل؛ الزائر يرى مودال تسجيل دخول/إنشاء حساب.', 'arabic', False, None),
    ('4.7.15 طلب الإسعاف', 'Side title', True, None),
    ('نموذج في الموقع العام → جدول ambulance_requests → المدير يرسل سيارة / يكمل / يلغي. تحديث حالة توفر سيارة الإسعاف تلقائياً.', 'arabic', False, None),
    ('4.7.16 تصدير السجل الطبي PDF', 'Side title', True, None),
    ('مكتبة DomPDF — يشمل: بيانات المريض، التشخيصات، الوصفات، الأشعة، المختبر. متاح للطبيب والمريض والمدير.', 'arabic', False, None),
    ('4.7.17 إدارة الطابور (Queue Management)', 'Side title', True, None),
    ('جدول queue_tickets: رقم يومي، قسم، طبيب، موعد، أولوية، حالة (انتظار → نداء → عند الطبيب → مكتمل).', 'arabic', False, None),
    ('لوحة الاستقبال (/queue): إصدار رقم، تسجيل حضور موعد، فلترة. لوحة الطبيب (/doctor/queue): نداء التالي، بدء كشف، إنهاء.', 'arabic', False, None),
    ('شاشة عرض TV (/queue/display/section/{id})، تتبع عام (/queue/track)، بث Pusher + polling كل 5 ثوانٍ، تقدير وقت الانتظار.', 'arabic', False, None),
    ('4.7.18 بيانات تجريبية (Database Seeders)', 'Side title', True, None),
    ('Seeder رئيسي: HmsFullDemoSeeder — يحمّل كل بيانات العرض التجريبية بترتيب صحيح.', 'arabic', False, None),
    ('الأوامر: php artisan db:seed  |  php artisan db:seed --class=HmsFullDemoSeeder  |  php artisan migrate:fresh --seed', 'arabic', False, None),
    ('يشمل: AdminTableSeeder، SectionTableSeeder، DoctorTableSeeder، PatientTableSeeder، AppointmentBookingSeeder، InvoiceDemoSeeder، ExtendedFeaturesSeeder، BlogSiteSettingSeeder، QueueTicketSeeder.', 'arabic', False, None),
    ('حسابات تجريبية: المدير admin@gmail.com / 123456789 — المريض patient@yahoo.com / 12345678', 'arabic', False, None),
]

TOC_EXTRA = [
    ('4.7.6 الموقع الإلكتروني العام', 'toc 3'),
    ('4.7.7 نظام الإشعارات والتواصل', 'toc 3'),
    ('4.7.8 الوصفات الإلكترونية (e-Prescription)', 'toc 3'),
    ('4.7.9 لوحة التقارير والإحصائيات', 'toc 3'),
    ('4.7.10 جدولة أوقات الطبيب', 'toc 3'),
    ('4.7.11 تكامل التأمين الكامل', 'toc 3'),
    ('4.7.12 تذكير المواعيد', 'toc 3'),
    ('4.7.13 تقييم الأطباء', 'toc 3'),
    ('4.7.14 إدارة المدونة', 'toc 3'),
    ('4.7.15 طلب الإسعاف', 'toc 3'),
    ('4.7.16 تصدير السجل الطبي PDF', 'toc 3'),
    ('4.7.17 إدارة الطابور (Queue Management)', 'toc 3'),
    ('4.7.18 بيانات تجريبية (Database Seeders)', 'toc 3'),
]


# ── fix steps ────────────────────────────────────────────────────────

def fix_cover_page(doc: Document) -> None:
    prep = find_paragraph(doc, equals='إعداد:')
    if prep:
        # Cover may use table cells — fill empty sibling paragraphs or insert new ones
        anchor = prep
        for name in STUDENTS:
            nxt_para = None
            nxt_el = anchor._p.getnext()
            if nxt_el is not None and nxt_el.tag.endswith('p'):
                nxt_para = Paragraph(nxt_el, anchor._parent)
            if nxt_para is not None and not nxt_para.text.strip():
                try:
                    set_paragraph(nxt_para, name, style='arabic',
                                  align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16)
                    anchor = nxt_para
                    continue
                except Exception:
                    pass
            anchor = insert_after(anchor, name, style='arabic')
            anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in anchor.runs:
                run.font.size = Pt(16)
                run.font.name = 'Traditional Arabic'

    for para in doc.paragraphs[:20]:
        t = para.text.strip()
        if 'تطوير نظام' in t and 'المستشفى' in t and len(t) < 60:
            try:
                set_paragraph(para, 'تطوير مستشفى الشام التخصصي', style='arabic', bold=True,
                              align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=22)
            except Exception:
                para.text = 'تطوير مستشفى الشام التخصصي'
            break

    subtitle = find_paragraph(doc, contains='مشروع تخرج')
    if subtitle:
        try:
            set_paragraph(subtitle, subtitle.text.replace('الالكترونيات', 'الإلكترونيات'),
                          style='arabic', align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16)
        except Exception:
            subtitle.text = subtitle.text.replace('الالكترونيات', 'الإلكترونيات')


def clean_toc_junk(doc: Document) -> None:
    """Remove body content wrongly inserted inside the table of contents."""
    toc_start = find_index(doc, equals='فهرس المحتويات')
    summary = find_index(doc, equals='ملخص المشروع')
    if toc_start is None or summary is None:
        return

    junk_prefixes = (
        '12.7.', '4.7.', 'insurance_id', '/reports', 'المسارات:', 'شاشة عرض',
        'جدول queue', '• ', 'حسابات تجريبية', '| ', 'Seeder', 'جدول Seeders',
        'php artisan', 'Seeder رئيسي', 'أو إعادة بناء', 'الأمر:',
    )
    for para in list(doc.paragraphs)[toc_start + 1:summary]:
        t = para.text.strip()
        style = para.style.name
        if style.startswith('toc'):
            if any(t.startswith(p) for p in junk_prefixes):
                delete_paragraph(para)
            elif t in ('73', '79'):
                delete_paragraph(para)
            elif 'في الفصل التالي' in t:
                delete_paragraph(para)
        elif style == 'Normal' and toc_start < find_index(doc, equals='ملخص المشروع'):
            if any(t.startswith(p) for p in junk_prefixes):
                delete_paragraph(para)


def normalize_toc_entries(doc: Document) -> None:
    mapping = {
        '12. الإجراء العملي': '4.1 المقدمة',
        '12.1 المقدمة': '4.1 المقدمة',
        '12.2 التقنيات المستخدمة': '4.2 التقنيات المستخدمة',
        '12.3 تصميم وتنفيذ قاعدة البيانات': '4.3 تصميم وتنفيذ قاعدة البيانات',
        '12.4 تطوير الواجهات البرمجية': '4.4 تطوير الواجهات البرمجية',
        '12.5 تطوير واجهات المستخدم': '4.5 تطوير واجهات المستخدم',
        '12.6 تنفيذ نظام المصادقة والتفويض': '4.6 تنفيذ نظام المصادقة والتفويض',
        '12.7 تنفيذ الوحدات الرئيسية': '4.7 تنفيذ الوحدات الرئيسية',
        '12.8 الاختبارات والتقييم': '4.8 الاختبارات والتقييم',
        '12.9 التحديات والحلول': '4.9 التحديات والحلول',
        '12.10 خاتمة الفصل': '4.10 خاتمة الفصل',
        '12. الفصل الخامس: الاختبار': 'الفصل الخامس — الاختبار',
        '12.1 أنواع الاختبارات': '5.1 أنواع الاختبارات',
        '12.1.1 اختبار الوحدة (Unit Testing)': '5.1.1 اختبار الوحدة (Unit Testing)',
        '12.1.2 اختبار التكامل (Integration Testing)': '5.1.2 اختبار التكامل (Integration Testing)',
        '12.1.3 اختبار النظام (System Testing)': '5.1.3 اختبار النظام (System Testing)',
        '12.1.4 اختبار القبول (Acceptance Testing)': '5.1.4 اختبار القبول (Acceptance Testing)',
        '12.2 سيناريوهات الاختبار': '5.2 سيناريوهات الاختبار',
        '12.3 نتائج الاختبار': '5.3 نتائج الاختبار',
        '12.4 خاتمة الفصل': '5.4 خاتمة الفصل',
        '13. الفصل السادس: النتائج والاستنتاجات': 'الفصل السادس — النتائج والاستنتاجات',
        '13.1 النتائج الرئيسية': '6.1 النتائج الرئيسية',
        '13.1.1 النتائج الوظيفية': '6.1.1 النتائج الوظيفية',
        '13.1.2 النتائج التقنية': '6.1.2 النتائج التقنية',
        '13.2 الإنجازات': '6.2 الإنجازات',
        '13.3 الاستنتاجات': '6.3 الاستنتاجات',
        '13.4 التوصيات': '6.4 التوصيات',
        '13.4.1 توصيات للتطوير المستقبلي': '6.4.1 توصيات للتطوير المستقبلي',
        '13.4.2 توصيات للاستخدام': '6.4.2 توصيات للاستخدام',
        '13.5 خاتمة الفصل': '6.5 خاتمة الفصل',
        '15. الخلاصة والاستنتاجات': 'الفصل السابع — الخلاصة والاستنتاجات',
        '16. الأفاق المستقبلية': 'الفصل الثامن — الأفاق المستقبلية',
        '17. المراجع': 'الفصل التاسع — المراجع',
    }

    toc_end = find_index(doc, equals='ملخص المشروع')
    if toc_end is None:
        return

    for para in list(doc.paragraphs)[:toc_end]:
        if not para.style.name.startswith('toc'):
            continue
        clean = strip_toc_page_num(para.text)
        for old, new in mapping.items():
            if clean == old or clean.startswith(old + '\t'):
                set_paragraph(para, new, style=para.style.name)
                break
        else:
            if '\t' in para.text:
                set_paragraph(para, strip_toc_page_num(para.text), style=para.style.name)

    # Insert missing advanced TOC entries before 4.8
    anchor = find_paragraph(doc, startswith='4.8')
    if anchor is None:
        anchor = find_paragraph(doc, startswith='12.8')
    if anchor:
        for text, style in reversed(TOC_EXTRA):
            if find_paragraph(doc, equals=text) is None:
                insert_before(anchor, text, style=style)


def renumber_body_sections(doc: Document) -> None:
    """Global section renumbering 12→4, 13→6, 15→7, 16→8, 17→9 in body text."""
    patterns = [
        (re.compile(r'\b12\.10\b'), '4.10'),
        (re.compile(r'\b12\.9\b'), '4.9'),
        (re.compile(r'\b12\.8\b'), '4.8'),
        (re.compile(r'\b12\.7\b'), '4.7'),
        (re.compile(r'\b12\.6\b'), '4.6'),
        (re.compile(r'\b12\.5\b'), '4.5'),
        (re.compile(r'\b12\.4\b'), '4.4'),
        (re.compile(r'\b12\.3\b'), '4.3'),
        (re.compile(r'\b12\.2\b'), '4.2'),
        (re.compile(r'\b12\.1\b'), '4.1'),
        (re.compile(r'\b13\.5\b'), '6.5'),
        (re.compile(r'\b13\.4\.2\b'), '6.4.2'),
        (re.compile(r'\b13\.4\.1\b'), '6.4.1'),
        (re.compile(r'\b13\.4\b'), '6.4'),
        (re.compile(r'\b13\.3\b'), '6.3'),
        (re.compile(r'\b13\.2\b'), '6.2'),
        (re.compile(r'\b13\.1\.2\b'), '6.1.2'),
        (re.compile(r'\b13\.1\.1\b'), '6.1.1'),
        (re.compile(r'\b13\.1\b'), '6.1'),
        (re.compile(r'\b15\.1\b'), '7.1'),
        (re.compile(r'\b16\.3\b'), '8.3'),
        (re.compile(r'\b16\.2\b'), '8.2'),
        (re.compile(r'\b16\.1\b'), '8.1'),
        (re.compile(r'\b16\.0\b'), '8.0'),
        (re.compile(r'\b17\.2\b'), '9.2'),
        (re.compile(r'\b17\.1\b'), '9.1'),
    ]

    chapter_headers = {
        '12. الإجراء العملي': '4. الإجراء العملي',
        '12. الفصل الخامس: الاختبار': 'الفصل الخامس — الاختبار',
        '13. الفصل السادس: النتائج والاستنتاجات': 'الفصل السادس — النتائج والاستنتاجات',
        '15. الخلاصة والاستنتاجات': 'الفصل السابع — الخلاصة والاستنتاجات',
        '16. الأفاق المستقبلية': 'الفصل الثامن — الأفاق المستقبلية',
        '17. المراجع': 'الفصل التاسع — المراجع',
    }

    summary_start = find_index(doc, equals='ملخص المشروع')
    if summary_start is None:
        return

    for para in doc.paragraphs[summary_start:]:
        if para.style.name.startswith('toc'):
            continue
        text = para.text
        for old, new in chapter_headers.items():
            text = text.replace(old, new)
        for rx, repl in patterns:
            text = re.sub(rx, repl, text)
        if text != para.text:
            bold = any(r.bold for r in para.runs) if para.runs else False
            set_paragraph(para, text, style=para.style.name, bold=bold)


def renumber_chapter5_testing(doc: Document) -> None:
    """Chapter 5 was also numbered 12.x — fix after chapter 4 renumber."""
    ch5 = find_index(doc, contains='الفصل الخامس')
    ch6 = find_index(doc, contains='الفصل السادس')
    if ch5 is None or ch6 is None:
        return

    mapping = {
        '4.1 ': '5.1 ',
        '4.1.1 ': '5.1.1 ',
        '4.1.2 ': '5.1.2 ',
        '4.1.3 ': '5.1.3 ',
        '4.1.4 ': '5.1.4 ',
        '4.2 ': '5.2 ',
        '4.2.1 ': '5.2.1 ',
        '4.2.2 ': '5.2.2 ',
        '4.2.3 ': '5.2.3 ',
        '4.2.4 ': '5.2.4 ',
        '4.2.5 ': '5.2.5 ',
        '4.3 ': '5.3 ',
        '4.3.1 ': '5.3.1 ',
        '4.3.2 ': '5.3.2 ',
        '4.3.3 ': '5.3.3 ',
        '4.4 ': '5.4 ',
    }
    for para in doc.paragraphs[ch5:ch6]:
        text = para.text
        if text.strip().startswith('4.10'):
            continue
        for old, new in mapping.items():
            if text.startswith(old):
                text = new + text[len(old):]
                break
        if text != para.text:
            set_paragraph(para, text, style=para.style.name)


def remove_duplicate_blocks(doc: Document) -> None:
    """Remove duplicate 11.4/11.5 and duplicate summary header."""
    start = find_index(doc, equals='11.4 التحديات والحلول')
    end = find_index(doc, contains='في الفصل التالي، سيتم عرض خطة الاختبار')
    if start is not None and end is not None and start < end:
        for para in list(doc.paragraphs)[start:end + 1]:
            delete_paragraph(para)

    # duplicate summary title
    seen_summary = False
    for para in list(doc.paragraphs):
        t = para.text.strip()
        if t == 'الخلاصة والاستنتاجات' and para.style.name == 'main title':
            if seen_summary:
                delete_paragraph(para)
            else:
                seen_summary = True


def rebuild_advanced_sections(doc: Document) -> None:
    """Remove reversed/misplaced 12.7.6+ block and insert clean ordered content."""
    start = find_index(doc, startswith='12.7.6')
    if start is None:
        start = find_index(doc, startswith='4.7.6')
    anchor = find_paragraph(doc, startswith='4.8')
    if anchor is None:
        anchor = find_paragraph(doc, startswith='12.8')

    if start is None or anchor is None:
        return

    end = find_index(doc, startswith='4.8')
    if end is None:
        end = find_index(doc, startswith='12.8')
    if end is None:
        return

    for para in list(doc.paragraphs)[start:end]:
        delete_paragraph(para)

    anchor = find_paragraph(doc, startswith='4.8') or find_paragraph(doc, startswith='12.8')
    for text, style, bold, align in reversed(SECTIONS_127_ADVANCED):
        insert_before(anchor, text, style=style, bold=bold)


def fix_summary_styles(doc: Document) -> None:
    idx = find_index(doc, equals='ملخص المشروع')
    ch1 = find_index(doc, equals='الفصل الأول')
    if idx is None or ch1 is None:
        return
    for para in doc.paragraphs[idx + 1:ch1]:
        t = para.text.strip()
        if not t:
            continue
        if para.style.name == 'Chapter title':
            set_paragraph(para, t, style='arabic')


def fix_chapter5_styles(doc: Document) -> None:
    ch5 = find_index(doc, contains='الفصل الخامس')
    ch6 = find_index(doc, contains='الفصل السادس')
    if ch5 is None or ch6 is None:
        return

    for i, para in enumerate(doc.paragraphs[ch5:ch6]):
        t = para.text.strip()
        if not t:
            continue
        if i == 0:
            set_paragraph(para, t, style='main title', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=18)
        elif t in ('Testing', 'Results and Conclusions'):
            set_paragraph(para, t, style='English', align=WD_ALIGN_PARAGRAPH.CENTER)
        elif re.match(r'^5\.\d', t) or re.match(r'^4\.\d', t):
            if re.match(r'^5\.\d+\.\d+', t) or t.startswith('5.2.') or t.startswith('5.3.'):
                set_paragraph(para, t, style='Sub-Side title', bold=True)
            elif re.match(r'^5\.\d+\s', t) or re.match(r'^5\.\d+$', t):
                set_paragraph(para, t, style='Side title', bold=True)
            else:
                set_paragraph(para, t, style='Sub-Side title', bold=True)
        elif para.style.name == 'Sub-Side title' and (t.startswith('-') or t.startswith('تم')):
            set_paragraph(para, t, style='arabic')


def fix_chapter3_numbering(doc: Document) -> None:
    ch3 = find_index(doc, equals='الفصل الثالث')
    ch4 = find_index(doc, equals='الفصل الرابع')
    if ch3 is None or ch4 is None:
        return
    fixes = {
        '10.2.2': '3-2-2',
        '10.2.3': '3-2-3',
        '10.3.1': '3-3-1',
        '10.3.2': '3-3-2',
        '10.3.3': '3-3-3',
        '10.4.1': '3-4-1',
        '10.4.2': '3-4-2',
        '10.4.3': '3-4-3',
        '10.5.1': '3-5-1',
        '10.5.2': '3-5-2',
        '9.5.2': '3-6-2',
    }
    for para in doc.paragraphs[ch3:ch4]:
        t = para.text.strip()
        for old, new in fixes.items():
            if t.startswith(old):
                set_paragraph(para, t.replace(old, new, 1), style=para.style.name, bold=True)
                break


def add_advanced_testing_scenarios(doc: Document) -> None:
    anchor = find_paragraph(doc, startswith='5.2.4') or find_paragraph(doc, startswith='4.2.4')
    if anchor is None:
        return
    if find_paragraph(doc, startswith='5.2.5'):
        return
    last = anchor
    lines = [
        ('5.2.5 اختبار الميزات المتقدمة', 'Sub-Side title', True),
        ('- عرض التقارير والرسوم البيانية للمدير', 'arabic', False),
        ('- إضافة وصفة إلكترونية مع التشخيص', 'arabic', False),
        ('- حجز موعد بتاريخ/وقت ضمن جدول الطبيب ومنع التعارض', 'arabic', False),
        ('- إنشاء مطالبة تأمين تلقائياً مع فاتورة مريض مؤمّن', 'arabic', False),
        ('- تشغيل php artisan appointments:send-reminders', 'arabic', False),
        ('- تقييم طبيب بعد موعد منتهي', 'arabic', False),
        ('- CRUD مقال في المدونة + إعجاب/تعليق (مودال للزائر)', 'arabic', False),
        ('- طلب إسعاف من الموقع وإرسال سيارة من المدير', 'arabic', False),
        ('- تصدير السجل الطبي PDF', 'arabic', False),
        ('- إصدار رقم طابور وتتبع الانتظار من /queue/track', 'arabic', False),
    ]
    for text, style, bold in lines:
        last = insert_after(last, text, style=style, bold=bold)


def fix_future_prospects(doc: Document) -> None:
    replace_all(doc, {
        'IoT و Blockchain و Queue Management.': 'IoT و Blockchain.',
        '5 | Queue Management | منفّذ (12.7.17)': '5 | Queue Management | منفّذ (4.7.17)',
    })
    p = find_paragraph(doc, startswith='8.0') or find_paragraph(doc, startswith='16.0')
    if p:
        set_paragraph(p, '8.0 الميزات المتقدمة المنفّذة', style='Side title', bold=True)
        nxt = p._p.getnext()
        if nxt is not None:
            para = Paragraph(nxt, p._parent)
            set_paragraph(
                para,
                'لوحة تقارير، e-Prescription، جدولة أطباء، تأمين كامل، تذكير مواعيد، '
                'تقييم أطباء، مدونة CRUD، طلب إسعاف، تصدير PDF، إدارة الطابور (Queue Management).',
                style='arabic',
            )


def apply_global_text_fixes(doc: Document) -> None:
    replace_all(doc, {
        'Livewire و Tailwind CSS': 'Livewire و Bootstrap (قالب Valex)',
        'Tailwind CSS': 'Bootstrap (قالب Valex)',
        'Blade Templates و Livewire و Tailwind CSS': 'Blade Templates و Livewire و Bootstrap (قالب Valex)',
        'Gates و Policies': 'Middleware متعدد الحراس (Multi-guard)',
        'Gates و Policies للتحكم في الوصول': 'Middleware متعدد الحراس (Multi-guard) للتحكم في الوصول',
        'Bootstrap (قالب Valex) هو إطار عمل CSS يتبع نهج Utility-First': 
            'Bootstrap (قالب Valex) هو قالب لوحة تحكم مبني على Bootstrap 5',
        '2. تحسين التقارير: إضافة تقارير أكثر تفصيلاً وإحصائيات متقدمة':
            '2. توسيع التقارير: تصدير Excel/PDF (تم تنفيذ لوحة Chart.js الأساسية)',
    })


def apply_paragraph_formatting(doc: Document) -> None:
    summary = find_index(doc, equals='ملخص المشروع')
    if summary is None:
        return

    for para in doc.paragraphs[summary:]:
        style = para.style.name
        if not para.text.strip():
            continue
        if style == 'arabic':
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.5
            pf.space_after = Pt(6)
            for run in para.runs:
                set_rtl(run)
                if not run.font.size:
                    run.font.size = Pt(14)
                if not run.font.name:
                    run.font.name = 'Traditional Arabic'
        elif style in ('Side title', 'Sub-Side title'):
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                set_rtl(run)
        elif style == 'main title':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_rtl(run)


def remove_empty_chapter_titles(doc: Document) -> None:
    for para in list(doc.paragraphs):
        if para.style.name == 'Chapter title' and not para.text.strip():
            delete_paragraph(para)


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    shutil.copy2(SOURCE, BACKUP)
    print(f'Backup: {BACKUP.name}')

    doc = Document(str(SOURCE))

    fix_cover_page(doc)
    apply_global_text_fixes(doc)
    clean_toc_junk(doc)
    normalize_toc_entries(doc)
    remove_duplicate_blocks(doc)
    rebuild_advanced_sections(doc)
    renumber_body_sections(doc)
    renumber_chapter5_testing(doc)
    fix_summary_styles(doc)
    fix_chapter5_styles(doc)
    fix_chapter3_numbering(doc)
    add_advanced_testing_scenarios(doc)
    fix_future_prospects(doc)
    remove_empty_chapter_titles(doc)
    apply_paragraph_formatting(doc)

    doc.save(str(OUTPUT))
    print(f'Saved: {OUTPUT.name}')
    print('Done — document formatted for university submission.')


if __name__ == '__main__':
    main()
