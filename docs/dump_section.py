# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document('تطوير نظام ادارة المستشفى - كامل.docx')

# Dump paragraphs 960-1140
for i in range(960, min(1145, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if t:
        print(f'{i:4d}| {p.style.name:20s}| {t[:85]}')
